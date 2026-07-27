"""Create a private, demonstrably sanitized Word template.

The approved reference is copied and neutralized in-place in the copy. Every
source-driven slot receives an ``AEP`` marker or a neutral image. The manifest
records a digest of those markers, the neutral media contract and the exact
capacity of the retained layout. The original reference is never modified.
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import json
import re
import shutil
import unicodedata
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml.ns import qn


SANITIZATION_CONTRACT_VERSION = 1
MARKER_PATTERN = re.compile(r"\{\{AEP_[A-Z0-9_]+\}\}")
NEUTRAL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAIAAAACCAYAAABytg0k"
    "AAAAFElEQVR4nGP8//8/AwgwgUkGBgYAMAYDAQGHVfcA"
    "AAAASUVORK5CYII="
)
TECHNICAL_HEADINGS = {
    "visao geral": "overview",
    "pontos positivos": "positive_points",
    "pontos criticos": "critical_points",
    "indicacoes de melhoria": "improvements",
}
TECHNICAL_BOUNDARY_PREFIXES = (
    "perguntas de maior relevancia",
    "priorizacoes recomendadas",
    "plano de acao geral integrado",
    "conclusao tecnica",
    "termo de encerramento",
    "assinaturas",
    "pagina institucional",
)


def _normalize(value: str) -> str:
    decomposed = unicodedata.normalize("NFKD", value or "")
    ascii_text = "".join(
        character for character in decomposed if not unicodedata.combining(character)
    )
    ascii_text = re.sub(r"[^a-zA-Z0-9]+", " ", ascii_text).strip().casefold()
    return " ".join(ascii_text.split())


def _marker(name: str) -> str:
    return f"{{{{AEP_{name}}}}}"


def _set_text(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _cell_text(cell, value: str) -> None:
    _set_text(cell.paragraphs[0], value)
    for paragraph in cell.paragraphs[1:]:
        _set_text(paragraph, "")


def _unique_cells(row) -> list[Any]:
    cells: list[Any] = []
    seen: set[Any] = set()
    for cell in row.cells:
        element = cell._tc
        if element not in seen:
            cells.append(cell)
            seen.add(element)
    return cells


def _hash_text(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()


def _mark_paragraph(
    paragraph,
    marker_name: str,
    source_text_hashes: set[str],
) -> None:
    original = " ".join(paragraph.text.split())
    if original and not MARKER_PATTERN.fullmatch(original):
        source_text_hashes.add(_hash_text(original))
    _set_text(paragraph, _marker(marker_name))


def _mark_cell(cell, marker_name: str, source_text_hashes: set[str]) -> None:
    original = " ".join(cell.text.split())
    if original and not MARKER_PATTERN.fullmatch(original):
        source_text_hashes.add(_hash_text(original))
    _cell_text(cell, _marker(marker_name))


def _iter_body_paragraphs(document) -> Iterable[Any]:
    yield from document.paragraphs
    for table in document.tables:
        seen: set[Any] = set()
        for row in table.rows:
            for cell in row.cells:
                element = cell._tc
                if element in seen:
                    continue
                seen.add(element)
                yield from cell.paragraphs


def _replace_literal(
    document,
    literal: str,
    replacement: str,
    source_text_hashes: set[str],
) -> None:
    if not literal or MARKER_PATTERN.fullmatch(literal):
        return
    for paragraph in _iter_body_paragraphs(document):
        if literal not in paragraph.text:
            continue
        original = " ".join(paragraph.text.split())
        if original:
            source_text_hashes.add(_hash_text(original))
        _set_text(paragraph, paragraph.text.replace(literal, replacement))


def _paragraph_image_parts(document) -> list[Any]:
    parts: list[Any] = []
    for paragraph in document.paragraphs:
        for blip in paragraph._p.xpath(".//a:blip"):
            relationship_id = blip.get(qn("r:embed"))
            if relationship_id and relationship_id in paragraph.part.related_parts:
                parts.append(paragraph.part.related_parts[relationship_id])
    return parts


def _marker_inventory(document) -> list[str]:
    markers: list[str] = []
    for paragraph in _iter_body_paragraphs(document):
        markers.extend(MARKER_PATTERN.findall(paragraph.text))
    return sorted(markers)


def _inventory_digest(markers: list[str]) -> str:
    payload = json.dumps(markers, ensure_ascii=True, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _find_paragraph_index(document, normalized_text: str) -> int | None:
    return next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if _normalize(paragraph.text) == normalized_text
        ),
        None,
    )


def _neutralize_revision_history(
    document,
    source_text_hashes: set[str],
) -> None:
    table = document.tables[0]
    for row_index, row in enumerate(table.rows[2:], start=2):
        for column_index, cell in enumerate(_unique_cells(row)):
            _mark_cell(
                cell,
                f"REVISION_R{row_index}_C{column_index}",
                source_text_hashes,
            )


def _neutralize_official_ghes(
    document,
    source_text_hashes: set[str],
) -> int:
    table = document.tables[1]
    data_rows = list(table.rows[1:])
    ghe_rows = [
        row
        for row in data_rows
        if re.search(r"(?i)\bGHE\s*[-_:.]?\s*0*\d+\b", row.cells[0].text)
    ]
    if not ghe_rows:
        # A valid retained template has one final TOTAL row.
        ghe_rows = data_rows[:-1]
    if not ghe_rows:
        raise ValueError("O template não contém slots oficiais de GHE.")
    for index, row in enumerate(ghe_rows, start=1):
        cells = _unique_cells(row)
        _mark_cell(cells[0], f"GHE_SLOT_{index}_LABEL", source_text_hashes)
        if len(cells) > 1:
            _mark_cell(cells[1], f"GHE_SLOT_{index}_POPULATION", source_text_hashes)
    total_row = next((row for row in data_rows if row not in ghe_rows), None)
    if total_row is not None:
        cells = _unique_cells(total_row)
        _cell_text(cells[0], "TOTAL")
        if len(cells) > 1:
            _mark_cell(cells[1], "TOTAL_POPULATION", source_text_hashes)
    return len(ghe_rows)


def _neutralize_diagnostic_summary(
    document,
    source_text_hashes: set[str],
) -> None:
    table = document.tables[2]
    marker_names = (
        "ERGO_BASE_DATE",
        "PSYCHOSOCIAL_BASE_DATE",
        "TOTAL_POPULATION_SUMMARY",
        "OFFICIAL_GHE_LIST",
    )
    for row_index, marker_name in enumerate(marker_names):
        if row_index >= len(table.rows):
            break
        cells = _unique_cells(table.rows[row_index])
        if cells:
            _mark_cell(cells[-1], marker_name, source_text_hashes)


def _neutralize_ergo_tables(
    document,
    capacity: int,
    source_text_hashes: set[str],
) -> None:
    first_table = 3
    for slot in range(capacity):
        base = first_table + slot * 4
        if base + 3 >= len(document.tables):
            raise ValueError("O template não contém todos os blocos Ergo declarados.")
        _, meta, counts, questions = document.tables[base : base + 4]
        meta_cells = _unique_cells(meta.rows[0])
        if len(meta_cells) < 4:
            raise ValueError("Um bloco Ergo possui metadados incompatíveis.")
        _mark_cell(meta_cells[1], f"ERGO_{slot + 1}_COMPANY", source_text_hashes)
        _mark_cell(meta_cells[3], f"ERGO_{slot + 1}_TITLE", source_text_hashes)
        count_cells = _unique_cells(counts.rows[0])
        if len(count_cells) < 3:
            raise ValueError("Um bloco Ergo possui contadores incompatíveis.")
        _mark_cell(count_cells[0], f"ERGO_{slot + 1}_YES", source_text_hashes)
        _mark_cell(count_cells[2], f"ERGO_{slot + 1}_NO", source_text_hashes)
        for row_index, row in enumerate(questions.rows[1:], start=1):
            for column_index, cell in enumerate(_unique_cells(row)):
                _mark_cell(
                    cell,
                    f"ERGO_{slot + 1}_QUESTION_R{row_index}_C{column_index}",
                    source_text_hashes,
                )


def _neutralize_technical_question_tables(
    document,
    capacity: int,
    source_text_hashes: set[str],
) -> None:
    first_table = 3 + capacity * 4
    for slot in range(capacity):
        table_index = first_table + slot
        if table_index >= len(document.tables):
            raise ValueError("O template não contém todas as tabelas técnicas.")
        for row_index, row in enumerate(
            document.tables[table_index].rows[1:],
            start=1,
        ):
            for column_index, cell in enumerate(_unique_cells(row)):
                _mark_cell(
                    cell,
                    f"TECH_QUESTION_{slot + 1}_R{row_index}_C{column_index}",
                    source_text_hashes,
                )


def _neutralize_priorities_and_actions(
    document,
    capacity: int,
    source_text_hashes: set[str],
) -> tuple[int, int]:
    priorities_index = 3 + capacity * 5
    actions_index = priorities_index + 1
    if actions_index >= len(document.tables):
        raise ValueError("O template não contém priorização e plano de ação.")
    for prefix, table in (
        ("PRIORITY", document.tables[priorities_index]),
        ("ACTION", document.tables[actions_index]),
    ):
        for row_index, row in enumerate(table.rows[1:], start=1):
            for column_index, cell in enumerate(_unique_cells(row)):
                _mark_cell(
                    cell,
                    f"{prefix}_R{row_index}_C{column_index}",
                    source_text_hashes,
                )
    return priorities_index, actions_index


def _paragraph_is_boundary(text: str) -> bool:
    key = _normalize(text)
    return key in TECHNICAL_HEADINGS or key.startswith(TECHNICAL_BOUNDARY_PREFIXES)


def _neutralize_technical_paragraphs(
    document,
    capacity: int,
    source_text_hashes: set[str],
) -> dict[str, Any]:
    paragraphs = document.paragraphs
    anchor = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if "resultado dos riscos psicossociais por ghe"
            in _normalize(paragraph.text)
        ),
        None,
    )
    if anchor is None:
        raise ValueError("A âncora das análises técnicas não foi encontrada.")
    terminal = next(
        (
            index
            for index, paragraph in enumerate(paragraphs[anchor + 1 :], anchor + 1)
            if _normalize(paragraph.text).startswith("priorizacoes recomendadas")
        ),
        len(paragraphs),
    )
    title_indexes = [
        index
        for index, paragraph in enumerate(paragraphs[anchor + 1 : terminal], anchor + 1)
        if len(paragraph.text.strip()) <= 160
        and re.match(r"(?i)^\s*GHE\s*[-_:.]?\s*0*\d+\b", paragraph.text)
    ]
    if len(title_indexes) != capacity:
        raise ValueError(
            "A quantidade de análises técnicas diverge da capacidade do template."
        )

    slots: list[dict[str, Any]] = []
    for slot, title_index in enumerate(title_indexes, start=1):
        region_end = (
            title_indexes[slot] if slot < len(title_indexes) else terminal
        )
        _mark_paragraph(
            paragraphs[title_index],
            f"TECH_GHE_SLOT_{slot}_TITLE",
            source_text_hashes,
        )
        section_slots: dict[str, dict[str, Any]] = {}
        for heading_index in range(title_index + 1, region_end):
            category = TECHNICAL_HEADINGS.get(
                _normalize(paragraphs[heading_index].text)
            )
            if category is None:
                continue
            body_indexes: list[int] = []
            for body_index in range(heading_index + 1, region_end):
                paragraph = paragraphs[body_index]
                if _paragraph_is_boundary(paragraph.text):
                    break
                if paragraph._p.xpath(".//w:drawing"):
                    continue
                if not paragraph.text.strip():
                    continue
                _mark_paragraph(
                    paragraph,
                    f"TECH_GHE_SLOT_{slot}_{category.upper()}_{len(body_indexes) + 1}",
                    source_text_hashes,
                )
                body_indexes.append(body_index)
            section_slots[category] = {
                "heading_index": heading_index,
                "body_indexes": body_indexes,
            }
        slots.append(
            {
                "title_index": title_index,
                "sections": section_slots,
            }
        )

    conclusion_heading = _find_paragraph_index(document, "conclusao tecnica")
    conclusion_body: list[int] = []
    if conclusion_heading is None:
        raise ValueError("A seção de conclusão técnica não foi encontrada.")
    for index in range(conclusion_heading + 1, len(paragraphs)):
        paragraph = paragraphs[index]
        key = _normalize(paragraph.text)
        if key.startswith(("termo de encerramento", "assinaturas", "pagina institucional")):
            break
        if paragraph._p.xpath(".//w:drawing") or not paragraph.text.strip():
            continue
        _mark_paragraph(
            paragraph,
            f"CONCLUSION_{len(conclusion_body) + 1}",
            source_text_hashes,
        )
        conclusion_body.append(index)
    return {
        "technical": slots,
        "conclusion": {
            "heading_index": conclusion_heading,
            "body_indexes": conclusion_body,
        },
    }


def _neutralize_dynamic_images(document, capacity: int) -> dict[str, Any]:
    parts = _paragraph_image_parts(document)
    dynamic_indices = [
        1,
        4,
        *range(5, 5 + capacity * 2),
    ]
    if not dynamic_indices or max(dynamic_indices) >= len(parts):
        raise ValueError("O template não contém todos os slots de imagem dinâmicos.")
    source_hashes: list[str] = []
    for index in dynamic_indices:
        part = parts[index]
        blob = bytes(part.blob)
        source_hashes.append(hashlib.sha256(blob).hexdigest())
        if not hasattr(part, "_blob"):
            raise ValueError("Um slot de imagem não pode ser neutralizado.")
        part._blob = NEUTRAL_PNG
    return {
        "dynamic_body_image_indices": dynamic_indices,
        "source_dynamic_media_sha256": sorted(set(source_hashes)),
        "neutral_media_sha256": hashlib.sha256(NEUTRAL_PNG).hexdigest(),
    }


def _scrub_properties(document) -> None:
    properties = document.core_properties
    properties.author = "Automatizador de Documentos AEP"
    properties.last_modified_by = "Automatizador de Documentos AEP"
    properties.title = "Template privado sanitizado AEP"
    properties.subject = "Template local sem dados empresariais de origem"
    properties.keywords = "AEP; template privado; sanitizado"
    properties.comments = (
        "Template privado sanitizado para compilação local. "
        "Não versionar nem distribuir."
    )


def _verify_sanitized_template(
    template_path: Path,
    markers: list[str],
    image_contract: dict[str, Any],
) -> None:
    document = Document(template_path)
    if _marker_inventory(document) != markers:
        raise ValueError("O inventário de marcadores mudou durante a gravação.")
    parts = _paragraph_image_parts(document)
    neutral_hash = image_contract["neutral_media_sha256"]
    for index in image_contract["dynamic_body_image_indices"]:
        if hashlib.sha256(parts[index].blob).hexdigest() != neutral_hash:
            raise ValueError("Um slot de imagem dinâmica não foi neutralizado.")
    all_media_hashes = {
        hashlib.sha256(part.blob).hexdigest() for part in parts
    }
    residual = (
        set(image_contract["source_dynamic_media_sha256"])
        - {neutral_hash}
    ) & all_media_hashes
    if residual:
        raise ValueError("Uma mídia dinâmica da referência permaneceu no template.")


def prepare(source: Path, destination: Path) -> tuple[Path, Path]:
    source = source.resolve()
    destination = destination.resolve()
    if source == destination:
        raise ValueError("O gabarito original e o template devem ser arquivos distintos.")
    if not source.is_file():
        raise FileNotFoundError("O gabarito Word não foi encontrado.")
    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)

    document = Document(destination)
    if len(document.tables) < 20:
        raise ValueError("O gabarito não possui a estrutura tabular esperada.")
    source_text_hashes: set[str] = set()

    identification_index = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if "identificacao da empresa" in _normalize(paragraph.text)
        ),
        None,
    )
    if identification_index is None:
        raise ValueError("A identificação da empresa não foi encontrada.")
    prior = [
        paragraph
        for paragraph in document.paragraphs[:identification_index]
        if paragraph.text.strip()
    ]
    if len(prior) < 2:
        raise ValueError("Os slots de capa não foram encontrados.")
    company_literal = prior[-2].text.strip()
    competence_literal = prior[-1].text.strip()
    _replace_literal(
        document,
        company_literal,
        _marker("COMPANY_NAME"),
        source_text_hashes,
    )
    _replace_literal(
        document,
        competence_literal,
        _marker("COMPETENCE"),
        source_text_hashes,
    )

    _neutralize_revision_history(document, source_text_hashes)
    capacity = _neutralize_official_ghes(document, source_text_hashes)
    _neutralize_diagnostic_summary(document, source_text_hashes)
    _neutralize_ergo_tables(document, capacity, source_text_hashes)
    _neutralize_technical_question_tables(
        document,
        capacity,
        source_text_hashes,
    )
    priorities_index, actions_index = _neutralize_priorities_and_actions(
        document,
        capacity,
        source_text_hashes,
    )
    paragraph_slots = _neutralize_technical_paragraphs(
        document,
        capacity,
        source_text_hashes,
    )
    image_contract = _neutralize_dynamic_images(document, capacity)
    _scrub_properties(document)
    document.save(destination)

    sanitized = Document(destination)
    markers = _marker_inventory(sanitized)
    if not markers:
        raise ValueError("O template sanitizado não contém marcadores auditáveis.")
    _verify_sanitized_template(destination, markers, image_contract)

    manifest = {
        "schema_version": 1,
        "created_at": datetime.now(UTC).isoformat(),
        "source_sha256": hashlib.sha256(source.read_bytes()).hexdigest(),
        "template_sha256": hashlib.sha256(destination.read_bytes()).hexdigest(),
        "paragraph_count": len(sanitized.paragraphs),
        "table_count": len(sanitized.tables),
        "inline_shape_count": len(sanitized.inline_shapes),
        "slots": {
            "body_image_order": {
                "brand": 0,
                "registration_card": 1,
                "methodology": 2,
                "methodology_flow": 3,
                "psychosocial_summary": 4,
                "psychosocial_ghe_pairs_start": 5,
            },
            "tables": {
                "revision_history": 0,
                "official_ghes": 1,
                "diagnostic_summary": 2,
                "ergo_first": 3,
                "technical_questions_first": 3 + capacity * 4,
                "priorities": priorities_index,
                "action_plan": actions_index,
            },
            "paragraphs": paragraph_slots,
        },
        "sanitization": {
            "contract_version": SANITIZATION_CONTRACT_VERSION,
            "status": "sanitized",
            "marker_count": len(markers),
            "marker_inventory_sha256": _inventory_digest(markers),
            "source_dynamic_text_sha256": sorted(source_text_hashes),
            **image_contract,
            "capacity": {
                "official_ghes": capacity,
                "ergo_blocks": capacity,
                "psychosocial_ghe_pairs": capacity,
                "technical_ghe_sections": capacity,
            },
        },
    }
    manifest_path = destination.with_suffix(".manifest.json")
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return destination, manifest_path


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path)
    args = parser.parse_args()
    template, manifest = prepare(args.source, args.destination)
    print(f"Template privado: {template}")
    print(f"Manifesto privado: {manifest}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
