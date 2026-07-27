"""Gera fixtures publicas, deterministicas e sem qualquer dado empresarial real.

Os artefatos produzidos existem apenas para testes automatizados e demonstracoes
locais do Automatizador de Documentos AEP. Todo conteudo e deliberadamente
ficticio e esta marcado como "SEM VALIDADE".
"""

from __future__ import annotations

import argparse
import base64
import hashlib
import io
import json
import math
import os
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageDraw, ImageFont


NOTICE = "DADOS 100% SINTETICOS - SEM VALIDADE"
SYNTHETIC_COMPANY = "EMPRESA SINTETICA HORIZONTE LTDA"
INVALID_CNPJ = "00.000.000/0000-00"
FIXED_TIME = datetime(2000, 1, 1, tzinfo=timezone.utc)
OUTPUT_FILES = (
    "ghe_sinteticos.xlsx",
    "ergo_sintetico.doc",
    "psicossocial_sintetico.docx",
    "tecnico_integrado_sintetico.docx",
    "cartao_cnpj_sintetico.png",
    "logo_sintetica_opcional.png",
    "template_aep_sintetico.docx",
    "template_manifesto_sintetico.json",
)

NAVY = "17365D"
BLUE = "2F75B5"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"
RED = "B42318"
GREEN = "1E7D4D"
AMBER = "B7791F"


@dataclass(frozen=True)
class SyntheticGHE:
    code: str
    name: str
    sectors: tuple[str, ...]
    roles: tuple[str, ...]
    population: int
    favorability: int
    risk_class: str

    @property
    def label(self) -> str:
        return f"GHE {self.code} - {self.name}"


OFFICIAL_GHES: tuple[SyntheticGHE, ...] = (
    SyntheticGHE(
        code="10",
        name="Operacao Experimental",
        sectors=("Unidade Alfa",),
        roles=("Operador de bancada", "Apoio de processo"),
        population=5,
        favorability=72,
        risk_class="Atencao",
    ),
    SyntheticGHE(
        code="20",
        name="Atendimento Simulado",
        sectors=("Unidade Beta", "Suporte remoto"),
        roles=("Analista de atendimento",),
        population=8,
        favorability=58,
        risk_class="Prioritario",
    ),
    SyntheticGHE(
        code="30",
        name="Laboratorio Didatico",
        sectors=("Unidade Gama",),
        roles=("Tecnico de laboratorio",),
        population=3,
        favorability=81,
        risk_class="Monitoramento",
    ),
)

EXTRA_ERGO_GHE = SyntheticGHE(
    code="99",
    name="Bloco Extra Sem Correspondencia",
    sectors=("Area temporaria",),
    roles=("Funcao demonstrativa",),
    population=2,
    favorability=45,
    risk_class="Nao aplicavel",
)


def parse_args() -> argparse.Namespace:
    project_root = Path(__file__).resolve().parents[1]
    default_output = project_root / "tests" / "fixtures" / "public_synthetic"
    parser = argparse.ArgumentParser(
        description="Gera fixtures publicas 100% sinteticas e sem validade."
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=default_output,
        help=f"Diretorio de destino (padrao: {default_output})",
    )
    return parser.parse_args()


def font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = (
        ("arialbd.ttf" if bold else "arial.ttf"),
        ("DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"),
        str(
            Path(os.environ.get("WINDIR", r"C:\Windows"))
            / "Fonts"
            / ("arialbd.ttf" if bold else "arial.ttf")
        ),
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"
        if bold
        else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    )
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_width(draw: ImageDraw.ImageDraw, text: str, selected_font: ImageFont.ImageFont) -> int:
    box = draw.textbbox((0, 0), text, font=selected_font)
    return box[2] - box[0]


def draw_centered(
    draw: ImageDraw.ImageDraw,
    y: int,
    text: str,
    selected_font: ImageFont.ImageFont,
    fill: str,
    width: int,
) -> None:
    x = max(0, (width - text_width(draw, text, selected_font)) // 2)
    draw.text((x, y), text, font=selected_font, fill=fill)


def wrap_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    selected_font: ImageFont.ImageFont,
    max_width: int,
) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if current and text_width(draw, candidate, selected_font) > max_width:
            lines.append(current)
            current = word
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def draw_notice_banner(draw: ImageDraw.ImageDraw, width: int, height: int) -> None:
    draw.rectangle((0, 0, width, height), fill="#FFF1F0")
    draw.line((0, height - 1, width, height - 1), fill=f"#{RED}", width=3)
    draw_centered(draw, 11, NOTICE, font(22, bold=True), f"#{RED}", width)


def create_psychosocial_image(
    output: Path,
    *,
    title: str,
    subtitle: str,
    kind: str,
    index: int,
) -> None:
    width, height = 1400, 820
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw_notice_banner(draw, width, 55)
    draw.text((70, 92), title, font=font(42, bold=True), fill=f"#{NAVY}")
    draw.text((72, 150), subtitle, font=font(24), fill=f"#{MID_GRAY}")

    chart_left, chart_top, chart_right, chart_bottom = 90, 220, 1310, 735
    draw.rounded_rectangle(
        (chart_left, chart_top, chart_right, chart_bottom),
        radius=24,
        fill="#F8FAFC",
        outline="#CBD5E1",
        width=3,
    )

    if kind == "panel":
        labels = ("Demanda", "Autonomia", "Apoio", "Clareza")
        values = (62 + index, 78 - index, 70 + index, 67)
        card_width = 260
        for position, (label, value) in enumerate(zip(labels, values)):
            x0 = 130 + position * 292
            draw.rounded_rectangle(
                (x0, 300, x0 + card_width, 610),
                radius=20,
                fill="#FFFFFF",
                outline="#D0D5DD",
                width=2,
            )
            draw.text((x0 + 24, 340), label, font=font(24, bold=True), fill=f"#{NAVY}")
            draw.text((x0 + 72, 430), f"{value}%", font=font(46, bold=True), fill=f"#{BLUE}")
            status = "SIMULADO"
            draw.text((x0 + 52, 535), status, font=font(20, bold=True), fill=f"#{GREEN}")
    elif kind == "bar":
        labels = ("Ritmo", "Postura", "Pausas", "Ambiente", "Comunicacao")
        values = (64, 78, 55, 71, 83)
        for position, (label, value) in enumerate(zip(labels, values)):
            y = 275 + position * 82
            draw.text((145, y), label, font=font(22, bold=True), fill=f"#{NAVY}")
            draw.rounded_rectangle((360, y, 1150, y + 34), radius=12, fill="#E4E7EC")
            draw.rounded_rectangle(
                (360, y, 360 + int(790 * value / 100), y + 34),
                radius=12,
                fill=f"#{BLUE}",
            )
            draw.text((1170, y - 1), f"{value}%", font=font(22, bold=True), fill=f"#{NAVY}")
    elif kind == "radar":
        center_x, center_y, radius = 700, 470, 205
        axes = 6
        values = (0.72, 0.58, 0.81, 0.66, 0.74, 0.63)
        for ring in range(1, 5):
            ring_points = []
            for axis in range(axes):
                angle = -math.pi / 2 + axis * 2 * math.pi / axes
                ring_points.append(
                    (
                        center_x + radius * ring / 4 * math.cos(angle),
                        center_y + radius * ring / 4 * math.sin(angle),
                    )
                )
            draw.polygon(ring_points, outline="#CBD5E1")
        data_points = []
        for axis, value in enumerate(values):
            angle = -math.pi / 2 + axis * 2 * math.pi / axes
            end = (
                center_x + radius * math.cos(angle),
                center_y + radius * math.sin(angle),
            )
            draw.line((center_x, center_y, *end), fill="#98A2B3", width=2)
            data_points.append(
                (
                    center_x + radius * value * math.cos(angle),
                    center_y + radius * value * math.sin(angle),
                )
            )
        draw.polygon(data_points, fill="#B9D7EE", outline=f"#{BLUE}")
        draw.line((*data_points[-1], *data_points[0]), fill=f"#{BLUE}", width=4)
        for x, y in data_points:
            draw.ellipse((x - 7, y - 7, x + 7, y + 7), fill=f"#{NAVY}")
    elif kind == "matrix":
        colors = ("#D1FADF", "#FEF0C7", "#FEDF89", "#FECDCA", "#FDA29B")
        cell = 82
        start_x, start_y = 455, 275
        for row in range(5):
            for column in range(5):
                color = colors[min(4, max(row, column))]
                x0 = start_x + column * cell
                y0 = start_y + (4 - row) * cell
                draw.rectangle(
                    (x0, y0, x0 + cell, y0 + cell),
                    fill=color,
                    outline="#FFFFFF",
                    width=3,
                )
                draw.text(
                    (x0 + 30, y0 + 25),
                    str((row + 1) * (column + 1)),
                    font=font(18, bold=True),
                    fill=f"#{NAVY}",
                )
        draw.text((515, 695), "Probabilidade (sintetica)", font=font(19, bold=True), fill=f"#{NAVY}")
        draw.text((235, 475), "Impacto", font=font(19, bold=True), fill=f"#{NAVY}")
    else:
        raise ValueError(f"Tipo visual desconhecido: {kind}")

    draw_centered(
        draw,
        775,
        "Artefato publico de teste - nao representa avaliacao tecnica",
        font(18),
        f"#{MID_GRAY}",
        width,
    )
    image.save(output, format="PNG", optimize=False)


def create_logo(output: Path) -> None:
    width, height = 900, 420
    image = Image.new("RGBA", (width, height), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((25, 25, 875, 395), radius=48, fill="#17365D", outline="#2F75B5", width=8)
    draw.ellipse((80, 100, 300, 320), fill="#D9EAF7")
    draw.polygon(((190, 125), (270, 285), (110, 285)), fill="#2F75B5")
    draw.text((340, 112), "HORIZONTE", font=font(54, bold=True), fill="#FFFFFF")
    draw.text((343, 184), "MARCA SINTETICA", font=font(30, bold=True), fill="#D9EAF7")
    draw.text((343, 245), "OPCIONAL - SEM VALIDADE", font=font(24, bold=True), fill="#FDA29B")
    image.save(output, format="PNG", optimize=False)


def create_cnpj_card(output: Path) -> None:
    width, height = 1400, 900
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw_notice_banner(draw, width, 62)
    draw.rectangle((60, 100, 1340, 840), outline="#344054", width=3)
    draw.text((105, 145), "COMPROVANTE SINTETICO DE INSCRICAO", font=font(35, bold=True), fill=f"#{NAVY}")
    draw.text((105, 205), "E DE SITUACAO CADASTRAL", font=font(35, bold=True), fill=f"#{NAVY}")
    draw.line((105, 270, 1295, 270), fill="#98A2B3", width=2)
    entries = (
        ("CNPJ SINTETICO", f"{INVALID_CNPJ} (NUMERO INVALIDO)"),
        ("RAZAO SOCIAL", SYNTHETIC_COMPANY),
        ("NOME FANTASIA", "HORIZONTE DEMONSTRACAO"),
        ("SITUACAO", "DOCUMENTO FICTICIO PARA TESTES"),
        ("EMISSAO", "01/01/2000"),
    )
    y = 315
    for label, value in entries:
        draw.text((110, y), label, font=font(20, bold=True), fill=f"#{MID_GRAY}")
        for line_number, line in enumerate(wrap_text(draw, value, font(29, bold=True), 1050)):
            draw.text(
                (110, y + 31 + line_number * 35),
                line,
                font=font(29, bold=True),
                fill=f"#{NAVY}",
            )
        y += 104
    draw.rounded_rectangle((100, 735, 1300, 815), radius=12, fill="#FFF1F0")
    draw_centered(
        draw,
        757,
        "NAO E DOCUMENTO FISCAL, CADASTRAL OU EMPRESARIAL",
        font(25, bold=True),
        f"#{RED}",
        width,
    )
    image.save(output, format="PNG", optimize=False)


def set_docx_core_properties(document: Document) -> None:
    properties = document.core_properties
    properties.title = NOTICE
    properties.subject = "Fixture publica do Automatizador de Documentos AEP"
    properties.author = "Gerador de fixtures sinteticas"
    properties.keywords = "sintetico; fixture; sem validade"
    properties.comments = "Nenhuma informacao deste arquivo representa pessoa ou empresa real."
    properties.created = FIXED_TIME
    properties.modified = FIXED_TIME
    properties.revision = 1


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Title", 24, NAVY, 0, 12),
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, NAVY, 10, 6),
        ("Heading 3", 11, NAVY, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = style_name != "Title"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(NOTICE)
    footer_run.bold = True
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(RED)


def add_notice(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(NOTICE)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(RED)


def style_table(table, *, header: bool = True, widths_cm: Sequence[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths_cm and column_index < len(widths_cm):
                cell.width = Cm(widths_cm[column_index])
            if header and row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if header and row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)


def add_table(document: Document, headers: Sequence[str], rows: Iterable[Sequence[str]], widths_cm: Sequence[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    style_table(table, widths_cm=widths_cm)


def add_page_break(document: Document) -> None:
    document.add_page_break()


def create_ghe_workbook(output: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "GHEs"
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "A2"

    headers = ("Codigo GHE", "Nome GHE", "Setores", "Cargos", "Quantidade", "Aviso")
    sheet.append(headers)
    for ghe in OFFICIAL_GHES:
        sheet.append(
            (
                ghe.code,
                ghe.name,
                "; ".join(ghe.sectors),
                "; ".join(ghe.roles),
                ghe.population,
                NOTICE,
            )
        )

    header_fill = PatternFill("solid", fgColor=NAVY)
    header_font = Font(name="Arial", size=11, bold=True, color=WHITE)
    body_font = Font(name="Arial", size=10, color="101828")
    light_side = Side(style="thin", color="D0D5DD")

    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = Border(bottom=Side(style="medium", color=BLUE))
    sheet.row_dimensions[1].height = 30

    for row in sheet.iter_rows(min_row=2, max_row=1 + len(OFFICIAL_GHES)):
        for cell in row:
            cell.font = body_font
            cell.alignment = Alignment(vertical="center", wrap_text=True)
            cell.border = Border(bottom=light_side)
        row[4].number_format = "#,##0"
        row[4].alignment = Alignment(horizontal="right", vertical="center")
        row[5].font = Font(name="Arial", size=9, bold=True, color=RED)
    for row_index in range(2, 2 + len(OFFICIAL_GHES)):
        sheet.row_dimensions[row_index].height = 36

    widths = (14, 29, 28, 34, 14, 38)
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.auto_filter.ref = f"A1:F{1 + len(OFFICIAL_GHES)}"

    summary = workbook.create_sheet("Resumo")
    summary.sheet_view.showGridLines = False
    summary["A1"] = NOTICE
    summary["A1"].font = Font(name="Arial", size=13, bold=True, color=RED)
    summary.merge_cells("A1:D1")
    summary["A3"] = "Indicador"
    summary["B3"] = "Valor"
    summary["A4"] = "Quantidade de GHEs"
    summary["B4"] = "=COUNTA('GHEs'!A2:A4)"
    summary["A5"] = "Populacao total"
    summary["B5"] = "=SUM('GHEs'!E2:E4)"
    for cell in summary[3]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
    for row in summary.iter_rows(min_row=4, max_row=5, max_col=2):
        for cell in row:
            cell.font = body_font
            cell.border = Border(bottom=light_side)
    summary.column_dimensions["A"].width = 30
    summary.column_dimensions["B"].width = 18

    workbook.properties.title = NOTICE
    workbook.properties.subject = "Fixture publica de GHEs"
    workbook.properties.creator = "Gerador de fixtures sinteticas"
    workbook.properties.description = "Nenhuma informacao representa pessoa ou empresa real."
    workbook.properties.keywords = "sintetico,fixture,sem validade"
    workbook.properties.created = FIXED_TIME.replace(tzinfo=None)
    workbook.properties.modified = FIXED_TIME.replace(tzinfo=None)
    workbook.calculation.fullCalcOnLoad = True
    workbook.calculation.forceFullCalc = True
    workbook.save(output)
    normalize_zip_metadata(output)


def encode_png_data_uri(ghe: SyntheticGHE) -> str:
    image = Image.new("RGB", (640, 180), "#F8FAFC")
    draw = ImageDraw.Draw(image)
    draw_notice_banner(draw, 640, 44)
    draw.text((28, 68), ghe.label, font=font(25, bold=True), fill=f"#{NAVY}")
    draw.text(
        (28, 113),
        "Diagrama ergonomico meramente ilustrativo",
        font=font(19),
        fill=f"#{MID_GRAY}",
    )
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=False)
    encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def create_ergo_html(output: Path) -> None:
    blocks: list[str] = []
    for position, ghe in enumerate((*OFFICIAL_GHES, EXTRA_ERGO_GHE), start=1):
        rows = (
            (
                "Pergunta:",
                f"Os recursos demonstrativos do bloco {ghe.code} estao ajustados ao trabalho simulado?",
            ),
            ("Resposta:", "Parcialmente atendido em cenario sintetico."),
            (
                "Observacao:",
                f"Observacao ficticia {position}; nao corresponde a inspecao, pessoa ou local real.",
            ),
            (
                "Orientacao:",
                "Orientacao ficticia: revisar o arranjo em teste e registrar a decisao aprovada.",
            ),
        )
        rows_html = "\n".join(
            f"<tr><th>{label}</th><td>{value}</td></tr>" for label, value in rows
        )
        extra_marker = (
            '<p class="warning">BLOCO EXTRA PARA TESTAR RECONCILIACAO: '
            "nao existe GHE oficial correspondente.</p>"
            if ghe.code == EXTRA_ERGO_GHE.code
            else ""
        )
        blocks.append(
            f"""
            <section class="ghe" data-ghe-code="{ghe.code}" data-order="{position}">
              <h2>{ghe.label}</h2>
              {extra_marker}
              <p><strong>Setores:</strong> {"; ".join(ghe.sectors)}</p>
              <p><strong>Cargos:</strong> {"; ".join(ghe.roles)}</p>
              <table>
                <thead><tr><th>Campo</th><th>Conteudo sintetico</th></tr></thead>
                <tbody>{rows_html}</tbody>
              </table>
              <p><strong>Pergunta:</strong> Ha alternancia postural prevista no exercicio ficticio?</p>
              <p><strong>Resposta:</strong> Sim, apenas para demonstracao de extracao.</p>
              <p><strong>Observacao:</strong> O texto preserva a ordem visual do HTML.</p>
              <p><strong>Orientacao:</strong> Nao usar este conteudo para decisao tecnica.</p>
              <img alt="Diagrama sintetico de {ghe.label}" src="{encode_png_data_uri(ghe)}">
            </section>
            """
        )

    html = f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="generator" content="Fixture publica do Automatizador AEP">
  <title>{NOTICE}</title>
  <style>
    body {{ font-family: Arial, sans-serif; color: #17365d; margin: 32px; }}
    .notice {{ color: #b42318; font-weight: bold; border: 2px solid #b42318; padding: 12px; }}
    .ghe {{ page-break-before: always; margin-top: 28px; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #98a2b3; padding: 8px; text-align: left; }}
    th {{ background: #d9eaf7; }}
    .warning {{ background: #fff1f0; color: #b42318; padding: 10px; font-weight: bold; }}
    img {{ max-width: 640px; height: auto; }}
  </style>
</head>
<body>
  <!-- {NOTICE}; arquivo HTML deliberadamente salvo com extensao .doc -->
  <h1>Relatorio Ergonomico Sintetico</h1>
  <p class="notice">{NOTICE}. Nao representa avaliacao ergonomica.</p>
  <p>Este arquivo testa a deteccao por conteudo: e HTML UTF-8, nao DOC binario OLE.</p>
  {"".join(blocks)}
</body>
</html>
"""
    output.write_text(html, encoding="utf-8", newline="\n")


def set_picture_alt_text(inline_shape, description: str) -> None:
    document_properties = inline_shape._inline.docPr
    document_properties.set("descr", f"{NOTICE} | {description}")
    document_properties.set("title", description)


def create_psychosocial_docx(output: Path) -> None:
    document = Document()
    configure_docx_styles(document)
    set_docx_core_properties(document)
    section = document.sections[0]
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)

    image_specs: list[tuple[str, str, str]] = [
        ("Painel geral", "Resumo visual de todos os GHEs sinteticos", "panel")
    ]
    for ghe in OFFICIAL_GHES:
        image_specs.extend(
            (
                (f"Painel | {ghe.label}", "Painel psicossocial sintetico", "panel"),
                (f"Grafico | {ghe.label}", "Resumo dos dominios sinteticos", "bar"),
                (f"Radar | {ghe.label}", "Radar psicossocial sintetico", "radar"),
                (f"Matriz de risco | {ghe.label}", "Matriz de risco sintetica", "matrix"),
            )
        )

    with tempfile.TemporaryDirectory(prefix="aep_fixture_psico_") as temporary:
        temporary_dir = Path(temporary)
        for index, (title, subtitle, kind) in enumerate(image_specs, start=1):
            image_path = temporary_dir / f"{index:02d}_{kind}.png"
            create_psychosocial_image(
                image_path,
                title=title,
                subtitle=subtitle,
                kind=kind,
                index=index,
            )
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run()
            inline_shape = run.add_picture(str(image_path), width=Cm(18.0))
            set_picture_alt_text(inline_shape, f"{title} | {subtitle}")
            if index < len(image_specs):
                run.add_break()
                add_page_break(document)

    document.save(output)
    normalize_zip_metadata(output)


def add_labeled_paragraph(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    value_run = paragraph.add_run(value)
    for run in (label_run, value_run):
        run.font.name = "Arial"
        run.font.size = Pt(10.5)


def create_integrated_technical_docx(output: Path) -> None:
    document = Document()
    configure_docx_styles(document)
    set_docx_core_properties(document)
    document.add_heading("Relatorio Tecnico Integrado Sintetico", 0)
    add_notice(document)
    document.add_paragraph(
        "Conteudo pre-aprovado exclusivamente para testar extracao literal. "
        "O compilador nao deve resumir, corrigir ou criar conclusoes a partir deste arquivo."
    )

    document.add_heading("Visao geral", level=1)
    document.add_paragraph(
        "A amostra ficticia contem tres GHEs oficiais e cenarios ergonomicos e psicossociais "
        "deliberadamente artificiais. Nenhuma afirmacao descreve ambiente, trabalhador ou empresa real."
    )

    for position, ghe in enumerate(OFFICIAL_GHES, start=1):
        document.add_heading(ghe.label, level=1)
        document.add_heading("Pontos positivos", level=2)
        document.add_paragraph(
            f"Texto sintetico aprovado {position}: existe um mecanismo demonstrativo de apoio "
            "e registro das rotinas neste GHE ficticio."
        )
        document.add_heading("Pontos criticos", level=2)
        document.add_paragraph(
            f"Texto sintetico aprovado {position}: a consistencia do ritmo simulado deve ser "
            "acompanhada sem extrapolar conclusoes."
        )
        document.add_heading("Indicacoes de melhoria", level=2)
        document.add_paragraph(
            "Manter o conteudo exatamente como fornecido: testar uma pausa didatica, registrar "
            "o resultado e submeter qualquer decisao a profissional habilitado."
        )
        document.add_heading("Perguntas de maior relevancia", level=2)
        add_table(
            document,
            ("Pergunta sintetica", "Favorabilidade", "Classe", "Leitura tecnica aprovada"),
            (
                (
                    "Os recursos de apoio do cenario ficticio estao disponiveis?",
                    f"{ghe.favorability}%",
                    ghe.risk_class,
                    "Leitura literal de teste; nao constitui avaliacao.",
                ),
                (
                    "O fluxo demonstrativo permite registrar intercorrencias?",
                    f"{max(10, ghe.favorability - 9)}%",
                    ghe.risk_class,
                    "Conteudo sintetico preservado sem reescrita.",
                ),
            ),
            (6.8, 2.8, 3.0, 5.0),
        )
        document.add_heading("Favorabilidade", level=2)
        add_labeled_paragraph(document, "Percentual favoravel", f"{ghe.favorability}%")
        document.add_heading("Classe / Classificacao", level=2)
        add_labeled_paragraph(document, "Classe", ghe.risk_class)
        document.add_heading("Leitura tecnica", level=2)
        document.add_paragraph(
            "Leitura tecnica sintetica e pre-aprovada: conservar o texto integral e vincula-lo "
            "somente ao GHE indicado no titulo."
        )

    document.add_heading("Priorizacao", level=1)
    add_table(
        document,
        ("Ordem", "GHE", "Prioridade", "Fundamentacao aprovada"),
        (
            ("1", OFFICIAL_GHES[1].label, "Alta", "Menor favorabilidade no conjunto sintetico."),
            ("2", OFFICIAL_GHES[0].label, "Media", "Acompanhamento demonstrativo programado."),
            ("3", OFFICIAL_GHES[2].label, "Baixa", "Manter monitoramento no cenario ficticio."),
        ),
        (1.5, 5.5, 2.5, 8.1),
    )

    document.add_heading("Plano de acao integrado", level=1)
    add_table(
        document,
        (
            "Acao",
            "Origem",
            "Prioridade",
            "Responsavel",
            "Prazo",
            "Indicador",
            "Evolucao/Registros",
        ),
        (
            (
                "Realizar oficina didatica de organizacao do fluxo.",
                "Ergo + psicossocial sinteticos",
                "Alta",
                "Papel funcional ficticio",
                "30 dias",
                "Registro da oficina",
                "",
            ),
            (
                "Revisar pausas do exercicio simulado.",
                "Ergo sintetico",
                "Media",
                "Papel funcional ficticio",
                "60 dias",
                "Checklist aprovado",
                "",
            ),
        ),
        (4.0, 3.0, 2.0, 3.0, 1.8, 2.8, 3.0),
    )
    document.add_paragraph(
        "O campo Evolucao/Registros foi mantido livre para preenchimento posterior.",
        style="Caption",
    )

    document.add_heading("Conclusao tecnica", level=1)
    document.add_paragraph(
        "Conclusao sintetica pre-aprovada: os dados ficticios demonstram o encadeamento entre "
        "fontes, reconciliacao e plano de acao. Este paragrafo deve ser compilado literalmente "
        "e nao possui validade tecnica, legal ou empresarial."
    )

    document.save(output)
    normalize_zip_metadata(output)


def add_template_slot(document: Document, title: str, marker: str, *, new_page: bool = False) -> None:
    if new_page:
        add_page_break(document)
    document.add_heading(title, level=1)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(marker)
    run.bold = True
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(BLUE)


def create_template_docx(output: Path) -> None:
    document = Document()
    configure_docx_styles(document)
    set_docx_core_properties(document)
    section = document.sections[0]
    section.different_first_page_header_footer = True

    document.add_heading("MODELO AEP SINTETICO", 0)
    add_notice(document)
    document.add_paragraph("{{RAZAO_SOCIAL}}", style="Subtitle")
    document.add_paragraph("Competencia: {{COMPETENCIA}}")
    document.add_paragraph("{{LOGO_EMPRESA_OPCIONAL}}")
    document.add_paragraph("{{CARTAO_CNPJ}}")

    slots = (
        ("Historico de revisoes", "{{HISTORICO_REVISOES}}"),
        ("Sumario", "{{SUMARIO}}"),
        ("Objetivo", "{{OBJETIVO}}"),
        ("Fundamentacao legal", "{{FUNDAMENTACAO_LEGAL}}"),
        ("Metodologia", "{{METODOLOGIA}}"),
        ("Hierarquia dos GHEs", "{{GHE_HIERARQUIA}}"),
        ("Resumo do diagnostico", "{{RESUMO_DIAGNOSTICO}}"),
        ("Resultados do Ergo", "{{RESULTADOS_ERGO}}"),
        ("Resumo psicossocial", "{{RESUMO_PSICOSSOCIAL}}"),
        ("Analises psicossociais por GHE", "{{ANALISES_PSICOSSOCIAIS}}"),
        ("Perguntas e favorabilidade", "{{PERGUNTAS_FAVORABILIDADE}}"),
        ("Priorizacoes", "{{PRIORIZACOES}}"),
        ("Plano de acao integrado", "{{PLANO_ACAO}}"),
        ("Conclusao tecnica", "{{CONCLUSAO_TECNICA}}"),
        ("Termo de encerramento", "{{TERMO_ENCERRAMENTO}}"),
        ("Assinaturas", "{{ASSINATURAS}}"),
        ("Pagina institucional", "{{PAGINA_INSTITUCIONAL}}"),
    )
    for index, (title, marker) in enumerate(slots):
        add_template_slot(document, title, marker, new_page=index in {0, 6, 11, 15})

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph(
        "Fim do modelo publico sintetico. Nenhum conteudo deste arquivo possui validade."
    )
    document.save(output)
    normalize_zip_metadata(output)


def create_template_manifest(output: Path) -> None:
    slots = [
        {
            "name": name,
            "marker": marker,
            "required": required,
            "content_type": content_type,
        }
        for name, marker, required, content_type in (
            ("razao_social", "{{RAZAO_SOCIAL}}", True, "text"),
            ("competencia", "{{COMPETENCIA}}", True, "text"),
            ("logo_empresa", "{{LOGO_EMPRESA_OPCIONAL}}", False, "image"),
            ("cartao_cnpj", "{{CARTAO_CNPJ}}", True, "image"),
            ("historico_revisoes", "{{HISTORICO_REVISOES}}", True, "table"),
            ("sumario", "{{SUMARIO}}", True, "field"),
            ("objetivo", "{{OBJETIVO}}", True, "rich_text"),
            ("fundamentacao_legal", "{{FUNDAMENTACAO_LEGAL}}", True, "rich_text"),
            ("metodologia", "{{METODOLOGIA}}", True, "rich_text"),
            ("ghe_hierarquia", "{{GHE_HIERARQUIA}}", True, "table"),
            ("resumo_diagnostico", "{{RESUMO_DIAGNOSTICO}}", True, "rich_text"),
            ("resultados_ergo", "{{RESULTADOS_ERGO}}", True, "rich_content"),
            ("resumo_psicossocial", "{{RESUMO_PSICOSSOCIAL}}", True, "image_group"),
            ("analises_psicossociais", "{{ANALISES_PSICOSSOCIAIS}}", True, "rich_content"),
            ("perguntas_favorabilidade", "{{PERGUNTAS_FAVORABILIDADE}}", True, "table"),
            ("priorizacoes", "{{PRIORIZACOES}}", True, "table"),
            ("plano_acao", "{{PLANO_ACAO}}", True, "table"),
            ("conclusao_tecnica", "{{CONCLUSAO_TECNICA}}", True, "rich_text"),
            ("termo_encerramento", "{{TERMO_ENCERRAMENTO}}", True, "rich_text"),
            ("assinaturas", "{{ASSINATURAS}}", True, "signature_area"),
            ("pagina_institucional", "{{PAGINA_INSTITUCIONAL}}", True, "rich_content"),
        )
    ]
    manifest = {
        "schema_version": 1,
        "fixture": True,
        "validity": "SEM VALIDADE",
        "privacy": "Conteudo 100% sintetico; nao representa pessoa ou empresa real.",
        "template": "template_aep_sintetico.docx",
        "page": {"size": "A4", "orientation": "portrait"},
        "slots": slots,
        "test_expectations": {
            "official_ghe_codes": [ghe.code for ghe in OFFICIAL_GHES],
            "official_ghe_count": len(OFFICIAL_GHES),
            "official_population": sum(ghe.population for ghe in OFFICIAL_GHES),
            "ergo_ghe_codes": [ghe.code for ghe in (*OFFICIAL_GHES, EXTRA_ERGO_GHE)],
            "unmatched_ergo_code": EXTRA_ERGO_GHE.code,
            "psychosocial_image_count": 1 + 4 * len(OFFICIAL_GHES),
            "logo_optional": True,
            "card_required": True,
            "action_log_column_must_be_blank": True,
        },
    }
    output.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def normalize_zip_metadata(path: Path) -> None:
    """Normaliza ordem e timestamps de DOCX/XLSX para reduzir diffs binarios."""
    temporary = path.with_suffix(path.suffix + ".normalized")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as target:
        for name in sorted(source.namelist()):
            source_info = source.getinfo(name)
            info = zipfile.ZipInfo(name, date_time=(2000, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = source_info.external_attr
            info.create_system = source_info.create_system
            target.writestr(info, source.read(name))
    temporary.replace(path)


def verify_xlsx(path: Path) -> dict[str, int]:
    workbook = load_workbook(path, read_only=True, data_only=False)
    sheet = workbook["GHEs"]
    headers = [cell.value for cell in next(sheet.iter_rows(min_row=1, max_row=1))]
    required = ["Codigo GHE", "Nome GHE", "Setores", "Cargos", "Quantidade"]
    if headers[:5] != required:
        raise RuntimeError(f"Cabecalhos inesperados na fixture XLSX: {headers[:5]}")
    rows = list(sheet.iter_rows(min_row=2, max_row=4, values_only=True))
    workbook.close()
    if len(rows) != 3:
        raise RuntimeError("A fixture XLSX deve conter exatamente tres GHEs oficiais.")
    total = sum(int(row[4]) for row in rows)
    return {"ghe_count": len(rows), "population": total}


def verify_docx(path: Path, *, expected_images: int | None = None) -> dict[str, int]:
    document = Document(path)
    with zipfile.ZipFile(path) as archive:
        images = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
    if expected_images is not None and len(images) != expected_images:
        raise RuntimeError(
            f"{path.name}: esperadas {expected_images} imagens; encontradas {len(images)}."
        )
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "images": len(images),
    }


def verify_outputs(output_dir: Path) -> dict[str, object]:
    missing = [name for name in OUTPUT_FILES if not (output_dir / name).is_file()]
    if missing:
        raise RuntimeError(f"Fixtures nao geradas: {', '.join(missing)}")

    workbook_result = verify_xlsx(output_dir / "ghe_sinteticos.xlsx")
    if workbook_result != {
        "ghe_count": len(OFFICIAL_GHES),
        "population": sum(ghe.population for ghe in OFFICIAL_GHES),
    }:
        raise RuntimeError(f"Resumo XLSX inesperado: {workbook_result}")

    ergo_bytes = (output_dir / "ergo_sintetico.doc").read_bytes()
    if not ergo_bytes.lstrip().lower().startswith(b"<!doctype html>"):
        raise RuntimeError("A fixture .doc deve ser HTML detectavel pelo conteudo.")
    ergo_text = ergo_bytes.decode("utf-8")
    if ergo_text.count('<section class="ghe"') != 4 or EXTRA_ERGO_GHE.label not in ergo_text:
        raise RuntimeError("A fixture Ergo deve conter quatro blocos, incluindo o extra.")

    psychosocial = verify_docx(
        output_dir / "psicossocial_sintetico.docx",
        expected_images=1 + 4 * len(OFFICIAL_GHES),
    )
    technical = verify_docx(output_dir / "tecnico_integrado_sintetico.docx")
    template = verify_docx(output_dir / "template_aep_sintetico.docx")
    if technical["tables"] < 5:
        raise RuntimeError("O relatorio tecnico sintetico nao contem as tabelas esperadas.")
    if template["paragraphs"] < 20:
        raise RuntimeError("O template sintetico nao contem todos os slots minimos.")

    manifest = json.loads(
        (output_dir / "template_manifesto_sintetico.json").read_text(encoding="utf-8")
    )
    if manifest.get("validity") != "SEM VALIDADE" or not manifest.get("fixture"):
        raise RuntimeError("Manifesto sintetico sem marcacao obrigatoria de validade.")

    hashes = {
        name: hashlib.sha256((output_dir / name).read_bytes()).hexdigest()
        for name in OUTPUT_FILES
    }
    return {
        "xlsx": workbook_result,
        "psychosocial": psychosocial,
        "technical": technical,
        "template": template,
        "sha256": hashes,
    }


def build_fixtures(output_dir: Path) -> dict[str, object]:
    output_dir.mkdir(parents=True, exist_ok=True)
    create_ghe_workbook(output_dir / "ghe_sinteticos.xlsx")
    create_ergo_html(output_dir / "ergo_sintetico.doc")
    create_psychosocial_docx(output_dir / "psicossocial_sintetico.docx")
    create_integrated_technical_docx(output_dir / "tecnico_integrado_sintetico.docx")
    create_cnpj_card(output_dir / "cartao_cnpj_sintetico.png")
    create_logo(output_dir / "logo_sintetica_opcional.png")
    create_template_docx(output_dir / "template_aep_sintetico.docx")
    create_template_manifest(output_dir / "template_manifesto_sintetico.json")
    return verify_outputs(output_dir)


def main() -> int:
    arguments = parse_args()
    output_dir = arguments.output_dir.expanduser().resolve()
    summary = build_fixtures(output_dir)
    print(f"Fixtures sinteticas geradas em: {output_dir}")
    print(
        "Resumo: "
        f"{summary['xlsx']['ghe_count']} GHEs, "
        f"populacao {summary['xlsx']['population']}, "
        f"{summary['psychosocial']['images']} imagens psicossociais."
    )
    print(NOTICE)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
