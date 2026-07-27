"""Structural and visual comparison for two locally rendered AEP documents."""

from __future__ import annotations

import argparse
import json
from dataclasses import asdict
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from PIL import Image, ImageChops, ImageStat

from app.services.document_renderer import inspect_docx_structure, render_docx


def _text(document) -> str:
    values = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text.strip() for cell in row.cells)
    return "\n".join(value for value in values if value)


def _headings(document) -> list[str]:
    result: list[str] = []
    for paragraph in document.paragraphs:
        style = (paragraph.style.name if paragraph.style else "").casefold()
        text = " ".join(paragraph.text.split())
        if text and (
            style.startswith(("heading", "título", "titulo"))
            or style == "list paragraph"
        ):
            result.append(text)
    return result


def _table_shapes(document) -> list[list[int]]:
    return [[len(table.rows), len(table.columns)] for table in document.tables]


def _page_metrics(reference: Path, candidate: Path) -> dict[str, float | int]:
    with Image.open(reference) as ref_image, Image.open(candidate) as new_image:
        reference_rgb = ref_image.convert("RGB")
        candidate_rgb = new_image.convert("RGB")
        if candidate_rgb.size != reference_rgb.size:
            candidate_rgb = candidate_rgb.resize(
                reference_rgb.size, Image.Resampling.LANCZOS
            )
        difference = ImageChops.difference(reference_rgb, candidate_rgb)
        mean = sum(ImageStat.Stat(difference).mean) / 3
        changed = difference.convert("L").point(
            lambda value: 255 if value > 16 else 0
        )
        histogram = changed.histogram()
        changed_pixels = sum(histogram[1:])
        total = reference_rgb.width * reference_rgb.height
        ref_gray = reference_rgb.convert("L")
        new_gray = candidate_rgb.convert("L")
        blank_reference = ImageStat.Stat(ref_gray).mean[0] > 252
        blank_candidate = ImageStat.Stat(new_gray).mean[0] > 252
        return {
            "mean_absolute_difference": round(mean, 3),
            "changed_pixel_ratio": round(changed_pixels / max(total, 1), 5),
            "reference_blank": int(blank_reference),
            "candidate_blank": int(blank_candidate),
        }


def compare(
    reference_path: Path,
    candidate_path: Path,
    work_dir: Path,
    *,
    libreoffice: Path | None = None,
    pdftoppm: Path | None = None,
) -> dict:
    work_dir.mkdir(parents=True, exist_ok=True)
    reference = Document(reference_path)
    candidate = Document(candidate_path)
    reference_text = _text(reference)
    candidate_text = _text(candidate)
    structural = {
        "reference": inspect_docx_structure(reference_path),
        "candidate": inspect_docx_structure(candidate_path),
        "reference_table_shapes": _table_shapes(reference),
        "candidate_table_shapes": _table_shapes(candidate),
        "reference_headings": _headings(reference),
        "candidate_headings": _headings(candidate),
        "text_similarity": round(
            SequenceMatcher(None, reference_text, candidate_text).ratio(), 5
        ),
        "candidate_contains_reference_text_ratio": round(
            sum(
                1
                for line in reference_text.splitlines()
                if line and line in candidate_text
            )
            / max(len(reference_text.splitlines()), 1),
            5,
        ),
    }
    reference_render = render_docx(
        reference_path,
        work_dir / "reference",
        libreoffice_path=libreoffice,
        pdftoppm_path=pdftoppm,
    )
    candidate_render = render_docx(
        candidate_path,
        work_dir / "candidate",
        libreoffice_path=libreoffice,
        pdftoppm_path=pdftoppm,
    )
    page_count = min(
        len(reference_render.page_images), len(candidate_render.page_images)
    )
    pages = [
        {
            "page": index + 1,
            **_page_metrics(
                reference_render.page_images[index],
                candidate_render.page_images[index],
            ),
        }
        for index in range(page_count)
    ]
    visual = {
        "reference_renderer": reference_render.renderer,
        "candidate_renderer": candidate_render.renderer,
        "reference_pages": len(reference_render.page_images),
        "candidate_pages": len(candidate_render.page_images),
        "compared_pages": page_count,
        "pages": pages,
        "warnings": reference_render.warnings + candidate_render.warnings,
    }
    return {"structural": structural, "visual": visual}


def _markdown(result: dict, reference: Path, candidate: Path) -> str:
    structural = result["structural"]
    visual = result["visual"]
    lines = [
        "# Comparação AEP",
        "",
        f"- Referência local: `{reference.name}`",
        f"- Documento automático: `{candidate.name}`",
        f"- Similaridade textual: {structural['text_similarity']:.2%}",
        (
            "- Linhas textuais da referência encontradas no automático: "
            f"{structural['candidate_contains_reference_text_ratio']:.2%}"
        ),
        (
            "- Páginas renderizadas: "
            f"{visual['reference_pages']} na referência e "
            f"{visual['candidate_pages']} no automático"
        ),
        (
            "- Renderizador: "
            f"{visual['reference_renderer']} / {visual['candidate_renderer']}"
        ),
        "",
        "## Comparação estrutural",
        "",
        "| Métrica | Referência | Automático |",
        "| --- | ---: | ---: |",
    ]
    for key, label in (
        ("sections", "Seções"),
        ("paragraphs", "Parágrafos"),
        ("tables", "Tabelas"),
        ("inline_shapes", "Imagens em linha"),
        ("media_parts", "Partes de mídia"),
    ):
        lines.append(
            f"| {label} | {structural['reference'][key]} | "
            f"{structural['candidate'][key]} |"
        )
    lines.extend(
        [
            "",
            "## Comparação visual por página",
            "",
            "| Página | Diferença média | Pixels alterados | Em branco inesperada |",
            "| ---: | ---: | ---: | --- |",
        ]
    )
    for page in visual["pages"]:
        blank = (
            "sim"
            if page["candidate_blank"] and not page["reference_blank"]
            else "não"
        )
        lines.append(
            f"| {page['page']} | {page['mean_absolute_difference']:.3f} | "
            f"{page['changed_pixel_ratio']:.2%} | {blank} |"
        )
    if visual["warnings"]:
        lines.extend(["", "## Avisos", ""])
        lines.extend(f"- {warning}" for warning in visual["warnings"])
    lines.extend(
        [
            "",
            "## Critério",
            "",
            (
                "A comparação mede estrutura, presença textual e rasterização. "
                "Ela não exige igualdade binária entre pacotes DOCX."
            ),
            "",
        ]
    )
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("reference", type=Path)
    parser.add_argument("candidate", type=Path)
    parser.add_argument("report", type=Path)
    parser.add_argument("--work-dir", type=Path)
    parser.add_argument("--libreoffice", type=Path)
    parser.add_argument("--pdftoppm", type=Path)
    args = parser.parse_args()
    work_dir = args.work_dir or args.report.parent / "comparison-render"
    result = compare(
        args.reference,
        args.candidate,
        work_dir,
        libreoffice=args.libreoffice,
        pdftoppm=args.pdftoppm,
    )
    args.report.parent.mkdir(parents=True, exist_ok=True)
    args.report.write_text(
        _markdown(result, args.reference, args.candidate), encoding="utf-8"
    )
    json_path = args.report.with_suffix(".metrics.json")
    json_path.write_text(
        json.dumps(result, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(args.report)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
