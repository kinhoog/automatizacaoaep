"""Execute an ignored local pilot from explicitly supplied private sources."""

from __future__ import annotations

import argparse
import hashlib
import json
from pathlib import Path

from docx import Document

from app.services.document_assembler import DocumentAssembler
from app.services.document_renderer import inspect_docx_structure
from app.services.normalization import ghe_name_similarity
from app.services.pipeline import DocumentPipeline
from app.services.reconciliation import apply_reconciliation_decisions
from app.services.validation import validate_normalized_aep


def _reference_identity(path: Path) -> tuple[str, str]:
    document = Document(path)
    marker = next(
        (
            index
            for index, paragraph in enumerate(document.paragraphs)
            if "identificação da empresa" in paragraph.text.casefold()
        ),
        None,
    )
    if marker is None:
        raise RuntimeError("A referência não contém o marcador de identificação.")
    prior = [
        paragraph.text.strip()
        for paragraph in document.paragraphs[:marker]
        if paragraph.text.strip()
    ]
    if len(prior) < 2:
        raise RuntimeError("Não foi possível localizar empresa e competência.")
    return prior[-2], prior[-1]


def _decisions(model) -> list[dict[str, object]]:
    decisions: list[dict[str, object]] = []
    for item in model.reconciliation.items:
        if item.status.value == "auto_matched":
            continue
        block = next(
            block for block in model.ergo.blocks if block.source_id == item.source_id
        )
        matches = sorted(
            (
                (ghe_name_similarity(block.source_name, ghe.name), ghe)
                for ghe in model.official_ghes
            ),
            key=lambda pair: pair[0],
            reverse=True,
        )
        if matches and matches[0][0] >= 0.95:
            decisions.append(
                {
                    "source_id": item.source_id,
                    "official_ghe_code": matches[0][1].canonical_code,
                    "reason": "Correspondência nominal exata confirmada no piloto local.",
                }
            )
        else:
            decisions.append(
                {
                    "source_id": item.source_id,
                    "not_applicable": True,
                    "reason": "Nenhum GHE oficial possui correspondência nominal suficiente.",
                }
            )
    return decisions


def _hash(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--ghe", required=True, type=Path)
    parser.add_argument("--psychosocial", required=True, type=Path)
    parser.add_argument("--ergo", required=True, type=Path)
    parser.add_argument("--technical", required=True, type=Path)
    parser.add_argument("--registration-card", required=True, type=Path)
    parser.add_argument("--logo", type=Path)
    parser.add_argument("--reference", required=True, type=Path)
    parser.add_argument("--template", required=True, type=Path)
    parser.add_argument("--manifest", type=Path)
    parser.add_argument("--output-document", required=True, type=Path)
    parser.add_argument("--validation-report", required=True, type=Path)
    parser.add_argument("--ergo-date", required=True)
    parser.add_argument("--psychosocial-date", required=True)
    parser.add_argument("--company-name")
    parser.add_argument("--competence")
    parser.add_argument("--compatibility", action="store_true")
    args = parser.parse_args()

    reference_company, reference_competence = _reference_identity(args.reference)
    files = {
        "ghe_spreadsheet": args.ghe,
        "psychosocial_report": args.psychosocial,
        "ergo_report": args.ergo,
        "integrated_report": args.technical,
        "cnpj_card": args.registration_card,
    }
    if args.logo:
        files["company_logo"] = args.logo
    model = DocumentPipeline().build_model(
        files,
        {
            "company_name": args.company_name or reference_company,
            "competence": args.competence or reference_competence,
            "ergo_reference_date": args.ergo_date,
            "psychosocial_reference_date": args.psychosocial_date,
            "analysis_mode": "integrated",
            "compatibility_mode": args.compatibility,
            "compatibility_acknowledged": args.compatibility,
        },
    )
    decisions = _decisions(model)
    model.reconciliation = apply_reconciliation_decisions(
        model.reconciliation, decisions, model.official_ghes
    )
    model.validation = validate_normalized_aep(model)
    if not model.validation.valid:
        codes = ", ".join(issue.code for issue in model.validation.errors)
        raise RuntimeError(f"Validação final reprovada: {codes}")

    args.output_document.parent.mkdir(parents=True, exist_ok=True)
    manifest_path = args.manifest or args.template.with_suffix(".manifest.json")
    DocumentAssembler(
        args.template,
        manifest_path if manifest_path.is_file() else None,
    ).assemble(model, args.output_document)
    report = {
        "schema_version": 1,
        "validation": model.validation.audit_dict(),
        "official_ghes": [ghe.audit_dict() for ghe in model.official_ghes],
        "total_population": model.total_population,
        "reconciliation": model.reconciliation.audit_dict(),
        "compatibility": (
            model.document.compatibility.audit_dict()
            if model.document.compatibility
            else None
        ),
        "output_structure": inspect_docx_structure(args.output_document),
        "source_integrity": {
            "ghe_sha256": _hash(args.ghe),
            "psychosocial_sha256": _hash(args.psychosocial),
            "ergo_sha256": _hash(args.ergo),
            "technical_sha256": _hash(args.technical),
            "registration_card_sha256": _hash(args.registration_card),
            "reference_sha256": _hash(args.reference),
        },
    }
    args.validation_report.write_text(
        json.dumps(report, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(args.output_document)
    print(args.validation_report)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
