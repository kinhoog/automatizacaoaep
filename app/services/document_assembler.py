"""Editable DOCX assembly from the normalized AEP model.

For the private production layout the assembler retains the approved Word
package and rewrites every variable source-driven slot.  A clean fallback
layout is available for public synthetic tests and installations that have not
yet prepared a private template.
"""

from __future__ import annotations

import copy
import base64
import hashlib
import io
import json
import re
from pathlib import Path
from typing import Any, Iterable, Mapping, Sequence

from docx import Document
from docx.document import Document as WordDocument
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models import (
    ActionPlanItem,
    ContentKind,
    ErgoBlock,
    GHE,
    ImageAsset,
    ImageRole,
    NormalizedAEP,
    ReconciliationStatus,
    TechnicalAnalysis,
    TechnicalCategory,
)
from app.services.image_processing import (
    asset_to_png_bytes,
    build_hierarchy_image,
    build_psychosocial_composite,
    match_psychosocial_block,
    trim_report_whitespace,
    image_to_png_bytes,
    open_asset,
)
from app.services.normalization import (
    canonical_ghe_code,
    clean_text,
    normalize_key,
)

_TECHNICAL_HEADINGS = {
    "visao geral": TechnicalCategory.OVERVIEW,
    "pontos positivos": TechnicalCategory.POSITIVE_POINTS,
    "pontos criticos": TechnicalCategory.CRITICAL_POINTS,
    "indicacoes de melhoria": TechnicalCategory.IMPROVEMENTS,
}
_TECHNICAL_SLOT_CATEGORIES = {
    "overview": TechnicalCategory.OVERVIEW,
    "positive_points": TechnicalCategory.POSITIVE_POINTS,
    "critical_points": TechnicalCategory.CRITICAL_POINTS,
    "improvements": TechnicalCategory.IMPROVEMENTS,
}

_DEFAULT_TABLE_SLOTS = {
    "revision_history": 0,
    "official_ghes": 1,
    "diagnostic_summary": 2,
    "ergo_first": 3,
    "technical_questions_first": 15,
    "priorities": 18,
    "action_plan": 19,
}

_DEFAULT_BODY_IMAGE_SLOTS = {
    "registration_card": 1,
    "psychosocial_summary": 4,
    "psychosocial_ghe_pairs_start": 5,
}

_SANITIZATION_CONTRACT_VERSION = 1
_TEMPLATE_MARKER_PATTERN = re.compile(r"\{\{AEP_[A-Z0-9_]+\}\}")
_NEUTRAL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAIAAAACCAYAAABytg0k"
    "AAAAFElEQVR4nGP8//8/AwgwgUkGBgYAMAYDAQGHVfcA"
    "AAAASUVORK5CYII="
)


class DocumentAssemblyError(RuntimeError):
    """Raised when a safe, editable output cannot be assembled."""


def _load_template_manifest(path: Path | None) -> dict[str, Any] | None:
    if path is None:
        raise ValueError("O manifesto sanitizado do template é obrigatório.")
    if not path.is_file():
        raise ValueError("O manifesto do template não foi encontrado.")
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict) or payload.get("schema_version") != 1:
        raise ValueError("O manifesto do template não é compatível.")
    return payload


def _manifest_slots(
    manifest: Mapping[str, Any],
) -> tuple[dict[str, int], dict[str, int]]:
    slots = manifest.get("slots")
    if not isinstance(slots, Mapping):
        raise ValueError("O manifesto não contém um mapa de slots.")
    table_source = slots.get("tables")
    image_source = slots.get("body_image_order")
    if not isinstance(table_source, Mapping) or not isinstance(
        image_source,
        Mapping,
    ):
        raise ValueError("O manifesto não declara todos os mapas de slots.")
    tables: dict[str, int] = {}
    images: dict[str, int] = {}
    for source, defaults, target in (
        (table_source, _DEFAULT_TABLE_SLOTS, tables),
        (image_source, _DEFAULT_BODY_IMAGE_SLOTS, images),
    ):
        for key in defaults:
            value = source.get(key)
            if not isinstance(value, int) or value < 0:
                raise ValueError("Um slot obrigatório do manifesto é inválido.")
            target[key] = value
    return tables, images


def _validate_template_manifest(
    document: WordDocument,
    template_path: Path,
    manifest: Mapping[str, Any] | None,
) -> None:
    if manifest is None:
        raise ValueError("O manifesto sanitizado do template é obrigatório.")
    expected_hash = str(manifest.get("template_sha256") or "").casefold()
    if not re.fullmatch(r"[0-9a-f]{64}", expected_hash):
        raise ValueError("O manifesto não identifica o template sanitizado.")
    if expected_hash:
        actual_hash = hashlib.sha256(template_path.read_bytes()).hexdigest()
        if actual_hash != expected_hash:
            raise ValueError("O template não corresponde ao manifesto de slots.")
    expected_counts = {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
    }
    for key, actual in expected_counts.items():
        expected = manifest.get(key)
        if not isinstance(expected, int) or expected != actual:
            raise ValueError("A estrutura do template diverge do manifesto.")

    sanitization = manifest.get("sanitization")
    if not isinstance(sanitization, Mapping):
        raise ValueError("O manifesto não comprova a sanitização do template.")
    if (
        sanitization.get("status") != "sanitized"
        or sanitization.get("contract_version")
        != _SANITIZATION_CONTRACT_VERSION
    ):
        raise ValueError("O contrato de sanitização do template é inválido.")
    markers = _template_marker_inventory(document)
    expected_count = sanitization.get("marker_count")
    expected_digest = str(
        sanitization.get("marker_inventory_sha256") or ""
    ).casefold()
    if (
        not isinstance(expected_count, int)
        or expected_count <= 0
        or len(markers) != expected_count
        or _template_marker_digest(markers) != expected_digest
    ):
        raise ValueError("Os slots neutralizados do template foram alterados.")
    dynamic_indices = sanitization.get("dynamic_body_image_indices")
    neutral_hash = str(sanitization.get("neutral_media_sha256") or "").casefold()
    source_text_hashes = sanitization.get("source_dynamic_text_sha256")
    source_media_hashes = sanitization.get("source_dynamic_media_sha256")

    def valid_hash_list(values: Any) -> bool:
        return (
            isinstance(values, list)
            and bool(values)
            and all(
                isinstance(value, str)
                and re.fullmatch(r"[0-9a-f]{64}", value.casefold())
                for value in values
            )
        )

    if (
        not isinstance(dynamic_indices, list)
        or not dynamic_indices
        or neutral_hash != hashlib.sha256(_NEUTRAL_PNG).hexdigest()
        or not all(isinstance(index, int) and index >= 0 for index in dynamic_indices)
        or not valid_hash_list(source_text_hashes)
        or not valid_hash_list(source_media_hashes)
    ):
        raise ValueError("O contrato de mídia neutralizada é inválido.")
    _, image_slots = _manifest_slots(manifest)
    capacity = _manifest_capacity(manifest)
    expected_dynamic_indices = {
        image_slots["registration_card"],
        image_slots["psychosocial_summary"],
        *range(
            image_slots["psychosocial_ghe_pairs_start"],
            image_slots["psychosocial_ghe_pairs_start"] + capacity * 2,
        ),
    }
    if set(dynamic_indices) != expected_dynamic_indices:
        raise ValueError("Os slots dinâmicos não estão todos neutralizados.")
    parts = _paragraph_image_parts(document)
    if max(dynamic_indices) >= len(parts):
        raise ValueError("O template perdeu um slot de imagem dinâmica.")
    if any(
        hashlib.sha256(parts[index].blob).hexdigest() != neutral_hash
        for index in dynamic_indices
    ):
        raise ValueError("Uma imagem dinâmica do template não está neutralizada.")


def _set_paragraph_text(paragraph: Paragraph, value: str) -> None:
    """Replace visible text while preserving paragraph and first-run styling."""

    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def _set_cell_text(cell, value: str) -> None:
    paragraph = cell.paragraphs[0]
    _set_paragraph_text(paragraph, value)
    for extra in cell.paragraphs[1:]:
        _set_paragraph_text(extra, "")


def _set_numbered_cell_text(cell, value: str) -> None:
    """Keep the template's number/body emphasis for an Ergo question."""

    paragraph = cell.paragraphs[0]
    match = re.match(r"^(\d+[.)]?)(\s*)(.*)$", value, flags=re.DOTALL)
    runs = paragraph.runs
    if not match or len(runs) < 2:
        _set_cell_text(cell, value)
        return

    number, separator, body = match.groups()
    body_run = next(
        (run for run in runs[1:] if run.bold is not True),
        runs[-1],
    )
    for run in runs:
        run.text = ""
    runs[0].text = number
    body_run.text = f"{separator or ' '}{body}"
    for extra in cell.paragraphs[1:]:
        _set_paragraph_text(extra, "")


def _set_ergo_count_cell(cell, count: int, label: str) -> None:
    """Retain the large counter and smaller status line from the template."""

    paragraphs = cell.paragraphs
    _set_paragraph_text(paragraphs[0], str(count))
    if len(paragraphs) == 1:
        status = cell.add_paragraph()
        status.alignment = paragraphs[0].alignment
        status.add_run(label)
    else:
        _set_paragraph_text(paragraphs[1], label)
    for extra in cell.paragraphs[2:]:
        _set_paragraph_text(extra, "")


def _clone_paragraph_after(reference: Paragraph, value: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    properties = reference._p.pPr
    if properties is not None:
        new_element.append(copy.deepcopy(properties))
    reference._p.addnext(new_element)
    paragraph = Paragraph(new_element, reference._parent)
    paragraph.add_run(value)
    return paragraph


def _remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._p.getparent()
    if parent is not None:
        parent.remove(paragraph._p)


def _paragraph_image_parts(document: WordDocument) -> list[object]:
    parts: list[object] = []
    for paragraph in document.paragraphs:
        for blip in paragraph._p.xpath(".//a:blip"):
            relationship_id = blip.get(qn("r:embed"))
            if relationship_id and relationship_id in paragraph.part.related_parts:
                parts.append(paragraph.part.related_parts[relationship_id])
    return parts


def _iter_template_paragraphs(document: WordDocument) -> Iterable[Paragraph]:
    yield from document.paragraphs
    for table in document.tables:
        seen: set[object] = set()
        for row in table.rows:
            for cell in row.cells:
                element = cell._tc
                if element in seen:
                    continue
                seen.add(element)
                yield from cell.paragraphs


def _template_marker_inventory(document: WordDocument) -> list[str]:
    markers: list[str] = []
    for paragraph in _iter_template_paragraphs(document):
        markers.extend(_TEMPLATE_MARKER_PATTERN.findall(paragraph.text))
    return sorted(markers)


def _template_marker_digest(markers: Sequence[str]) -> str:
    payload = json.dumps(
        list(markers),
        ensure_ascii=True,
        separators=(",", ":"),
    )
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _manifest_capacity(manifest: Mapping[str, Any]) -> int:
    sanitization = manifest.get("sanitization")
    capacity = (
        sanitization.get("capacity")
        if isinstance(sanitization, Mapping)
        else None
    )
    official = (
        capacity.get("official_ghes") if isinstance(capacity, Mapping) else None
    )
    expected = (
        official,
        capacity.get("ergo_blocks") if isinstance(capacity, Mapping) else None,
        (
            capacity.get("psychosocial_ghe_pairs")
            if isinstance(capacity, Mapping)
            else None
        ),
        (
            capacity.get("technical_ghe_sections")
            if isinstance(capacity, Mapping)
            else None
        ),
    )
    if (
        not isinstance(official, int)
        or official <= 0
        or any(value != official for value in expected)
    ):
        raise ValueError("A capacidade declarada pelo template é inválida.")
    return official


def _replace_scalar_markers(
    document: WordDocument,
    model: NormalizedAEP,
) -> None:
    values = {
        "{{AEP_COMPANY_NAME}}": model.company.legal_name,
        "{{AEP_COMPETENCE}}": model.document.competence,
    }
    for paragraph in _iter_template_paragraphs(document):
        value = paragraph.text
        for marker, replacement in values.items():
            value = value.replace(marker, replacement)
        if value != paragraph.text:
            _set_paragraph_text(paragraph, value)


def _assert_no_template_markers(document: WordDocument) -> None:
    residual = _template_marker_inventory(document)
    if residual:
        raise ValueError("O documento ainda contém slots privados não preenchidos.")


def _manifest_paragraph_slots(
    manifest: Mapping[str, Any],
) -> Mapping[str, Any]:
    slots = manifest.get("slots")
    paragraphs = slots.get("paragraphs") if isinstance(slots, Mapping) else None
    if not isinstance(paragraphs, Mapping):
        raise ValueError("O manifesto não contém os slots técnicos sanitizados.")
    return paragraphs


def _slot_paragraphs(
    paragraphs: Sequence[Paragraph],
    indexes: Any,
) -> list[Paragraph]:
    if not isinstance(indexes, list) or not all(
        isinstance(index, int) and 0 <= index < len(paragraphs)
        for index in indexes
    ):
        raise ValueError("Um slot de parágrafo do manifesto é inválido.")
    return [paragraphs[index] for index in indexes]


def _write_paragraph_slot(
    candidates: Sequence[Paragraph],
    values: Sequence[str],
    anchor: Paragraph,
) -> None:
    cleaned = [clean_text(value) for value in values if clean_text(value)]
    for index, paragraph in enumerate(candidates):
        if index < len(cleaned):
            _set_paragraph_text(paragraph, cleaned[index])
        else:
            _remove_paragraph(paragraph)
    if len(cleaned) <= len(candidates):
        return
    reference = candidates[-1] if candidates else anchor
    for value in cleaned[len(candidates) :]:
        reference = _clone_paragraph_after(reference, value)


def _replace_image_part(parts: list[object], index: int, png: bytes | None) -> None:
    if png is None or index >= len(parts):
        return
    part = parts[index]
    if hasattr(part, "_blob"):
        part._blob = png


def _find_paragraph(
    document: WordDocument,
    text: str,
    *,
    start: int = 0,
    exact: bool = False,
) -> tuple[int, Paragraph] | None:
    needle = normalize_key(text)
    for index, paragraph in enumerate(document.paragraphs[start:], start):
        haystack = normalize_key(paragraph.text)
        if (exact and haystack == needle) or (not exact and needle in haystack):
            return index, paragraph
    return None


def _add_update_fields_setting(document: WordDocument) -> None:
    settings = document.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _repair_psychosocial_toc_bookmark(document: WordDocument) -> None:
    """Repair the known retained-template PAGEREF target if it is absent."""

    bookmark_name = "_Toc235539110"
    existing = document.element.xpath(
        f".//w:bookmarkStart[@w:name='{bookmark_name}']"
    )
    if existing:
        return
    found = _find_paragraph(document, "Aplicação Formulário Psicossocial")
    if not found:
        return
    paragraph = found[1]
    used_ids = {
        int(item.get(qn("w:id")))
        for item in document.element.xpath(".//w:bookmarkStart")
        if (item.get(qn("w:id")) or "").isdigit()
    }
    bookmark_id = str(max(used_ids, default=0) + 1)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _normalize_psychosocial_page_break(document: WordDocument) -> None:
    """Remove the template break that can spill onto an otherwise blank page."""

    found = _find_paragraph(document, "Aplicação Formulário Psicossocial")
    if not found:
        return
    index, heading = found
    paragraphs = document.paragraphs
    for paragraph in paragraphs[max(0, index - 4) : index]:
        if clean_text(paragraph.text):
            continue
        for line_break in list(paragraph._p.xpath(".//w:br[@w:type='page']")):
            parent = line_break.getparent()
            if parent is not None:
                parent.remove(line_break)
    # The preceding Ergo table already fills the prior page in the populated
    # document. Forcing another break on the heading makes Word emit a
    # header/footer-only page when the template's blank separator spills.
    heading.paragraph_format.page_break_before = None


def _resize_table_rows(table: Table, desired: int) -> None:
    while len(table.rows) < desired:
        table.add_row()
    while len(table.rows) > desired:
        table._tbl.remove(table.rows[-1]._tr)


def _write_rows(
    table: Table,
    rows: Sequence[Sequence[str]],
    *,
    preserve_header: bool = False,
    preserve_numbered_first_column: bool = False,
) -> None:
    start = 1 if preserve_header else 0
    _resize_table_rows(table, start + len(rows))
    for row_index, values in enumerate(rows, start):
        cells = []
        seen_cells: set[object] = set()
        for cell in table.rows[row_index].cells:
            element = cell._tc
            if element not in seen_cells:
                cells.append(cell)
                seen_cells.add(element)
        for column, value in enumerate(values[: len(cells)]):
            cleaned = clean_text(value)
            if preserve_numbered_first_column and column == 0:
                _set_numbered_cell_text(cells[column], cleaned)
            else:
                _set_cell_text(cells[column], cleaned)
        for column in range(len(values), len(cells)):
            _set_cell_text(cells[column], "")


def _select_ergo_blocks(model: NormalizedAEP) -> list[ErgoBlock]:
    compatibility = model.document.compatibility
    by_id = {block.source_id: block for block in model.ergo.blocks}
    if compatibility and compatibility.included_ergo_source_ids:
        return [
            by_id[source_id]
            for source_id in compatibility.included_ergo_source_ids
            if source_id in by_id
        ]
    decision_by_id = {
        item.source_id: item
        for item in model.reconciliation.items
        if item.status
        in {
            ReconciliationStatus.AUTO_MATCHED,
            ReconciliationStatus.CONFIRMED,
        }
    }
    selected = [
        block for block in model.ergo.blocks if block.source_id in decision_by_id
    ]
    official_order = {
        ghe.canonical_code: index for index, ghe in enumerate(model.official_ghes)
    }
    return sorted(
        selected,
        key=lambda block: official_order.get(
            canonical_ghe_code(
                decision_by_id[block.source_id].official_ghe_code
            )
            or "",
            10_000 + block.order,
        ),
    )


def _ergo_question_rows(block: ErgoBlock) -> list[list[str]]:
    candidates: list[list[str]] = []
    for element in block.elements:
        if element.kind != ContentKind.TABLE:
            continue
        if element.rows:
            first = [clean_text(value) for value in element.rows[0]]
            if (
                first
                and re.match(r"^\d+\s*Sim\b", first[0], flags=re.IGNORECASE)
                and any(
                    re.match(r"^\d+\s*N[aã]o\b", value, flags=re.IGNORECASE)
                    for value in first[1:]
                )
            ):
                continue
        for row in element.rows:
            values = [clean_text(value) for value in row]
            key = normalize_key(" ".join(values))
            if not any(values) or key.startswith(
                ("pergunta resposta", "sim conforme", "empresa setor ghe")
            ):
                continue
            if len(values) == 2 and values[0].isdigit():
                answer_match = re.search(
                    r"(?:[✅☑✓❌✗]\s*)?(Sim|Não|Nao)\s*$",
                    values[1],
                    flags=re.IGNORECASE,
                )
                if answer_match:
                    question = values[1][: answer_match.start()].strip()
                    candidates.append(
                        [
                            f"{values[0]}. {question}",
                            answer_match.group(1).replace("Nao", "Não"),
                            "",
                        ]
                    )
                    continue
            if len(values) >= 2 and (
                re.match(r"^\d+(?:[.) -]|$)", values[0])
                or normalize_key(values[1]) in {"sim", "nao", "conforme"}
            ):
                candidates.append(
                    [
                        values[0],
                        values[1],
                        values[2] if len(values) > 2 else "",
                    ]
                )
    if candidates:
        return candidates
    rows: list[list[str]] = []
    length = max(len(block.questions), len(block.answers))
    for index in range(length):
        rows.append(
            [
                block.questions[index] if index < len(block.questions) else "",
                block.answers[index] if index < len(block.answers) else "",
                (
                    block.observations[index]
                    if index < len(block.observations)
                    else (
                        block.guidance[index]
                        if index < len(block.guidance)
                        else ""
                    )
                ),
            ]
        )
    return rows


def _technical_values(
    analysis: TechnicalAnalysis,
    category: TechnicalCategory,
) -> list[str]:
    values: list[str] = []
    for section in analysis.sections:
        if section.category == category:
            values.extend(clean_text(item) for item in section.paragraphs if item)
    if category == TechnicalCategory.TECHNICAL_READING:
        values.extend(analysis.technical_reading)
    return [value for value in values if value]


def _replace_body_after_heading(
    document: WordDocument,
    heading_index: int,
    values: Sequence[str],
) -> None:
    paragraphs = document.paragraphs
    candidates: list[Paragraph] = []
    for paragraph in paragraphs[heading_index + 1 :]:
        key = normalize_key(paragraph.text)
        if paragraph._p.xpath("./w:pPr/w:sectPr"):
            break
        if key in _TECHNICAL_HEADINGS or key.startswith(
            (
                "perguntas de maior relevancia",
                "ghe ",
                "priorizacoes recomendadas",
                "plano de acao geral integrado",
                "conclusao tecnica",
                "termo de encerramento",
                "assinaturas",
                "pagina institucional",
            )
        ):
            break
        if paragraph._p.xpath(".//w:drawing"):
            continue
        if clean_text(paragraph.text):
            candidates.append(paragraph)
    if not values:
        return
    for index, value in enumerate(values):
        if index < len(candidates):
            _set_paragraph_text(candidates[index], value)
        else:
            reference = candidates[-1] if candidates else paragraphs[heading_index]
            candidates.append(_clone_paragraph_after(reference, value))
    for paragraph in candidates[len(values) :]:
        _remove_paragraph(paragraph)


def _analysis_for_ghe(
    analyses: Sequence[TechnicalAnalysis],
    ghe: GHE,
) -> TechnicalAnalysis | None:
    for analysis in analyses:
        if canonical_ghe_code(
            analysis.official_ghe_code or analysis.ghe_code_hint
        ) == ghe.canonical_code:
            return analysis
    name = normalize_key(ghe.name)
    return next(
        (
            analysis
            for analysis in analyses
            if name and name in normalize_key(analysis.ghe_name_hint)
        ),
        None,
    )


def _populate_cover(document: WordDocument, model: NormalizedAEP) -> None:
    identification = _find_paragraph(document, "Identificação da empresa")
    if identification:
        paragraphs = document.paragraphs
        prior = [
            (index, paragraph)
            for index, paragraph in enumerate(paragraphs[: identification[0]])
            if clean_text(paragraph.text)
        ]
        if len(prior) >= 2:
            _set_paragraph_text(prior[-2][1], model.company.legal_name)
            _set_paragraph_text(prior[-1][1], model.document.competence)
        if model.company.logo:
            blank = next(
                (
                    paragraph
                    for paragraph in paragraphs[
                        prior[-2][0] + 1 : identification[0]
                    ]
                    if not paragraph.text and not paragraph._p.xpath(".//w:drawing")
                ),
                None,
            )
            if blank:
                blank.alignment = WD_ALIGN_PARAGRAPH.CENTER
                blank.add_run().add_picture(
                    io.BytesIO(asset_to_png_bytes(model.company.logo)),
                    width=Inches(1.25),
                )


def _populate_summary_tables(
    document: WordDocument,
    model: NormalizedAEP,
    table_slots: Mapping[str, int],
) -> None:
    required = max(
        table_slots["revision_history"],
        table_slots["official_ghes"],
        table_slots["diagnostic_summary"],
    )
    if required >= len(document.tables):
        raise ValueError("O template não contém as tabelas de resumo esperadas.")
    revision = document.tables[table_slots["revision_history"]]
    if len(revision.rows) <= 2:
        raise ValueError("O histórico de revisões não possui uma linha editável.")
    for row in revision.rows[2:]:
        seen_cells: set[object] = set()
        for cell in row.cells:
            element = cell._tc
            if element in seen_cells:
                continue
            seen_cells.add(element)
            _set_cell_text(cell, "")
    revision_cells = revision.rows[2].cells
    if len(revision_cells) >= 3:
        _set_cell_text(revision_cells[0], "0")
        _set_cell_text(revision_cells[1], model.document.competence)
        _set_cell_text(revision_cells[2], "Emissão inicial do documento")

    hierarchy = document.tables[table_slots["official_ghes"]]
    rows = [
        [f"{ghe.canonical_code} - {ghe.name}", str(ghe.population)]
        for ghe in model.official_ghes
    ]
    rows.append(["TOTAL", str(model.total_population)])
    _write_rows(hierarchy, rows, preserve_header=True)

    summary = document.tables[table_slots["diagnostic_summary"]]
    if len(summary.rows) >= 3:
        _set_cell_text(
            summary.rows[0].cells[-1],
            f"Formulário ERGO – levantamento de {model.document.ergo_base_date}",
        )
        _set_cell_text(
            summary.rows[1].cells[-1],
            "Formulário HSE – relatório emitido em "
            f"{model.document.psychosocial_base_date}",
        )
        _set_cell_text(
            summary.rows[2].cells[-1],
            f"{model.total_population} colaboradores",
        )
    if len(summary.rows) >= 4:
        _set_cell_text(
            summary.rows[3].cells[-1],
            ", ".join(ghe.canonical_code for ghe in model.official_ghes),
        )


def _populate_ergo(
    document: WordDocument,
    model: NormalizedAEP,
    table_slots: Mapping[str, int],
    capacity: int,
) -> None:
    first_table = table_slots["ergo_first"]
    if first_table + capacity * 4 > len(document.tables):
        raise ValueError("O template não contém todos os slots Ergo declarados.")
    blocks = _select_ergo_blocks(model)
    if len(blocks) > capacity:
        raise ValueError(
            "A quantidade de blocos Ergo excede a capacidade do template."
        )
    for index in range(capacity):
        base = first_table + index * 4
        if base + 3 >= len(document.tables):
            break
        title_table, meta_table, count_table, question_table = document.tables[
            base : base + 4
        ]
        if index >= len(blocks):
            for table in (title_table, meta_table, count_table, question_table):
                for row in table.rows:
                    for cell in row.cells:
                        _set_cell_text(cell, "")
            continue
        block = blocks[index]
        rows = _ergo_question_rows(block)
        answers = [normalize_key(row[1]) for row in rows if len(row) > 1]
        yes = sum(answer.startswith("sim") for answer in answers)
        no = sum(answer.startswith("nao") for answer in answers)
        if len(meta_table.rows[0].cells) >= 4:
            _set_cell_text(meta_table.rows[0].cells[1], model.company.legal_name)
            display_title = re.sub(
                r"^(?:setor\s*/\s*ghe\s*:?\s*)+",
                "",
                block.title,
                flags=re.IGNORECASE,
            )
            _set_cell_text(meta_table.rows[0].cells[3], display_title)
        if len(count_table.rows[0].cells) >= 3:
            _set_ergo_count_cell(
                count_table.rows[0].cells[0],
                yes,
                "Sim (Conforme)",
            )
            _set_ergo_count_cell(
                count_table.rows[0].cells[2],
                no,
                "Não (Não Conforme)",
            )
        _write_rows(
            question_table,
            rows,
            preserve_header=True,
            preserve_numbered_first_column=True,
        )


def _populate_psychosocial_images(
    document: WordDocument,
    model: NormalizedAEP,
    image_slots: Mapping[str, int],
    capacity: int,
) -> None:
    parts = _paragraph_image_parts(document)
    required = max(
        image_slots["registration_card"],
        image_slots["psychosocial_summary"],
        image_slots["psychosocial_ghe_pairs_start"] + capacity * 2 - 1,
    )
    if required >= len(parts):
        raise ValueError("O template não contém todos os slots de imagem esperados.")
    _replace_image_part(
        parts,
        image_slots["registration_card"],
        _NEUTRAL_PNG,
    )
    _replace_image_part(
        parts,
        image_slots["psychosocial_summary"],
        _NEUTRAL_PNG,
    )
    start = image_slots["psychosocial_ghe_pairs_start"]
    for index in range(capacity):
        _replace_image_part(parts, start + index * 2, _NEUTRAL_PNG)
        _replace_image_part(parts, start + index * 2 + 1, _NEUTRAL_PNG)
    if model.company.registration_card:
        _replace_image_part(
            parts,
            image_slots["registration_card"],
            asset_to_png_bytes(model.company.registration_card),
        )
    global_blocks = [
        block
        for block in model.psychosocial.blocks
        if not block.ghe_code_hint and not block.official_ghe_code
    ]
    global_images = (
        global_blocks[0].images if global_blocks else model.psychosocial.images
    )
    global_matrix = next(
        (
            image
            for image in global_images
            if image.role == ImageRole.RISK_MATRIX
        ),
        None,
    )
    if global_matrix is None:
        global_matrix = next(
            (
                image
                for image in reversed(global_images)
                if image.width_px
                and image.height_px
                and image.width_px / image.height_px >= 3
            ),
            None,
        )
    if global_matrix:
        _replace_image_part(
            parts,
            image_slots["psychosocial_summary"],
            image_to_png_bytes(
                trim_report_whitespace(open_asset(global_matrix))
            ),
        )

    for index, ghe in enumerate(model.official_ghes):
        block = match_psychosocial_block(model.psychosocial.blocks, ghe)
        if not block:
            continue
        composite, matrix = build_psychosocial_composite(block)
        _replace_image_part(parts, start + index * 2, composite)
        _replace_image_part(parts, start + index * 2 + 1, matrix)


def _populate_technical_text(
    document: WordDocument,
    model: NormalizedAEP,
    manifest: Mapping[str, Any],
    capacity: int,
) -> None:
    paragraph_slots = _manifest_paragraph_slots(manifest)
    technical_slots = paragraph_slots.get("technical")
    conclusion_slot = paragraph_slots.get("conclusion")
    if (
        not isinstance(technical_slots, list)
        or len(technical_slots) != capacity
        or not isinstance(conclusion_slot, Mapping)
    ):
        raise ValueError("Os slots técnicos do manifesto são incompatíveis.")
    paragraphs = list(document.paragraphs)
    for index, slot in enumerate(technical_slots):
        if not isinstance(slot, Mapping):
            raise ValueError("Um slot técnico do manifesto é inválido.")
        title_index = slot.get("title_index")
        sections = slot.get("sections")
        if (
            not isinstance(title_index, int)
            or not 0 <= title_index < len(paragraphs)
            or not isinstance(sections, Mapping)
        ):
            raise ValueError("Um slot técnico do manifesto é inválido.")
        ghe = model.official_ghes[index] if index < len(model.official_ghes) else None
        analysis = (
            _analysis_for_ghe(model.technical.analyses, ghe) if ghe else None
        )
        _set_paragraph_text(
            paragraphs[title_index],
            f"{ghe.canonical_code} — {ghe.name.upper()}" if ghe else "",
        )
        for slot_name, category in _TECHNICAL_SLOT_CATEGORIES.items():
            section = sections.get(slot_name)
            if not isinstance(section, Mapping):
                raise ValueError("Um corpo técnico do manifesto está ausente.")
            heading_index = section.get("heading_index")
            if (
                not isinstance(heading_index, int)
                or not 0 <= heading_index < len(paragraphs)
            ):
                raise ValueError("Uma âncora técnica do manifesto é inválida.")
            candidates = _slot_paragraphs(
                paragraphs,
                section.get("body_indexes"),
            )
            _write_paragraph_slot(
                candidates,
                _technical_values(analysis, category) if analysis else [],
                paragraphs[heading_index],
            )

    conclusion_heading = conclusion_slot.get("heading_index")
    if (
        not isinstance(conclusion_heading, int)
        or not 0 <= conclusion_heading < len(paragraphs)
    ):
        raise ValueError("A âncora da conclusão no manifesto é inválida.")
    _write_paragraph_slot(
        _slot_paragraphs(paragraphs, conclusion_slot.get("body_indexes")),
        model.technical.conclusion,
        paragraphs[conclusion_heading],
    )


def _populate_question_tables(
    document: WordDocument,
    model: NormalizedAEP,
    table_slots: Mapping[str, int],
    capacity: int,
) -> None:
    first_table = table_slots["technical_questions_first"]
    if first_table + capacity > len(document.tables):
        raise ValueError(
            "O template não contém todos os slots de perguntas técnicas."
        )
    for index in range(capacity):
        ghe = (
            model.official_ghes[index]
            if index < len(model.official_ghes)
            else None
        )
        analysis = (
            _analysis_for_ghe(model.technical.analyses, ghe) if ghe else None
        )
        rows: list[list[str]] = []
        if analysis:
            for section in analysis.sections:
                if section.category in {
                    TechnicalCategory.RELEVANT_QUESTIONS,
                    TechnicalCategory.FAVORABILITY,
                    TechnicalCategory.TECHNICAL_READING,
                }:
                    for source_table in section.tables:
                        rows.extend(source_table)
        normalized_rows = [
            row
            for row in rows
            if normalize_key(" ".join(row)).strip()
            and not normalize_key(" ".join(row)).startswith(
                "pergunta favoravel"
            )
        ]
        _write_rows(
            document.tables[first_table + index],
            normalized_rows,
            preserve_header=True,
        )


def _populate_priorities_and_actions(
    document: WordDocument,
    model: NormalizedAEP,
    table_slots: Mapping[str, int],
) -> None:
    priorities_index = table_slots["priorities"]
    actions_index = table_slots["action_plan"]
    if max(priorities_index, actions_index) >= len(document.tables):
        raise ValueError(
            "O template não contém os slots de priorização e plano de ação."
        )
    priority_rows = [
        [
            str(item.order + 1),
            item.ghe_code_hint or "",
            item.level or "",
            item.text,
        ]
        for item in model.technical.priorities
    ]
    _write_rows(
        document.tables[priorities_index],
        priority_rows,
        preserve_header=True,
    )
    action_rows = [
        [
            item.priority or "",
            item.ghe_code_hint or "",
            item.action,
            item.responsible or "",
            "",
        ]
        for item in model.technical.action_plan
    ]
    _write_rows(
        document.tables[actions_index],
        action_rows,
        preserve_header=True,
    )


def _set_document_metadata(document: WordDocument, model: NormalizedAEP) -> None:
    properties = document.core_properties
    properties.title = "Avaliação Ergonômica Preliminar"
    properties.subject = "Documento compilado localmente a partir das fontes validadas"
    properties.comments = (
        "Gerado pelo Automatizador de Documentos AEP. "
        "Conteúdo técnico preservado das fontes aprovadas."
    )
    _add_update_fields_setting(document)
    _normalize_psychosocial_page_break(document)
    _repair_psychosocial_toc_bookmark(document)


def _build_fallback_document(model: NormalizedAEP) -> WordDocument:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Title"].font.name = "Arial"
    styles["Title"].font.color.rgb = RGBColor(12, 79, 151)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AEP – Avaliação Ergonômica Preliminar")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(model.company.legal_name).bold = True
    if model.company.logo:
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo.add_run().add_picture(
            io.BytesIO(asset_to_png_bytes(model.company.logo)),
            width=Inches(1.3),
        )
    document.add_paragraph(model.document.competence).alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    if model.company.registration_card:
        card = document.add_paragraph()
        card.alignment = WD_ALIGN_PARAGRAPH.CENTER
        card.add_run().add_picture(
            io.BytesIO(asset_to_png_bytes(model.company.registration_card)),
            width=Inches(4.7),
        )
    document.add_page_break()

    ordered_sections = (
        "Histórico de revisões",
        "Sumário",
        "Objetivo",
        "Fundamentação legal",
        "Metodologia",
        "Hierarquia dos GHEs",
        "Resumo do diagnóstico",
        "Resultados do Ergo",
        "Resumo psicossocial",
        "Análise psicossocial por GHE",
        "Perguntas e favorabilidade",
        "Priorizações",
        "Plano de ação integrado",
        "Conclusão técnica",
        "Termo de encerramento",
        "Assinaturas",
        "Página institucional",
    )
    for heading in ordered_sections:
        document.add_heading(heading, level=1)
        if heading == "Hierarquia dos GHEs":
            image = document.add_paragraph()
            image.add_run().add_picture(
                io.BytesIO(
                    build_hierarchy_image(
                        model.company.legal_name, model.official_ghes
                    )
                ),
                width=Inches(6.4),
            )
            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "GHE"
            table.rows[0].cells[1].text = "Quantidade"
            for ghe in model.official_ghes:
                cells = table.add_row().cells
                cells[0].text = f"{ghe.canonical_code} — {ghe.name}"
                cells[1].text = str(ghe.population)
        elif heading == "Resultados do Ergo":
            for block in _select_ergo_blocks(model):
                document.add_heading(block.title, level=2)
                rows = _ergo_question_rows(block)
                if rows:
                    table = document.add_table(rows=1, cols=3)
                    table.style = "Table Grid"
                    for index, value in enumerate(
                        ("Pergunta", "Resposta", "Observação / Orientação")
                    ):
                        table.rows[0].cells[index].text = value
                    for row in rows:
                        cells = table.add_row().cells
                        for index, value in enumerate(row[:3]):
                            cells[index].text = value
        elif heading == "Análise psicossocial por GHE":
            for ghe in model.official_ghes:
                document.add_heading(
                    f"{ghe.canonical_code} — {ghe.name}", level=2
                )
                analysis = _analysis_for_ghe(model.technical.analyses, ghe)
                if analysis:
                    for section in analysis.sections:
                        document.add_heading(section.title, level=3)
                        for value in section.paragraphs:
                            document.add_paragraph(value)
        elif heading == "Plano de ação integrado":
            if model.technical.action_plan:
                table = document.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                headers = (
                    "Prioridade",
                    "GHE/Abrangência",
                    "Plano de ação",
                    "Responsável",
                    "Evolução/Registros",
                )
                for index, value in enumerate(headers):
                    table.rows[0].cells[index].text = value
                for item in model.technical.action_plan:
                    cells = table.add_row().cells
                    values = (
                        item.priority or "",
                        item.ghe_code_hint or "",
                        item.action,
                        item.responsible or "",
                        "",
                    )
                    for index, value in enumerate(values):
                        cells[index].text = value
        elif heading == "Conclusão técnica":
            for paragraph in model.technical.conclusion:
                document.add_paragraph(paragraph)
    return document


class DocumentAssembler:
    """Compile a normalized model into a real, editable Word document."""

    def __init__(
        self,
        template_path: str | Path | None = None,
        manifest_path: str | Path | None = None,
    ):
        self.template_path = Path(template_path) if template_path else None
        self.manifest_path = Path(manifest_path) if manifest_path else None

    def assemble(
        self,
        model: NormalizedAEP,
        output_path: str | Path,
    ) -> Path:
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        try:
            if self.template_path is not None:
                if not self.template_path.is_file():
                    raise ValueError("O template privado não foi encontrado.")
                document = Document(self.template_path)
                manifest = _load_template_manifest(self.manifest_path)
                _validate_template_manifest(
                    document,
                    self.template_path,
                    manifest,
                )
                table_slots, image_slots = _manifest_slots(manifest)
                capacity = _manifest_capacity(manifest)
                if len(model.official_ghes) > capacity:
                    raise ValueError(
                        "A quantidade de GHEs oficiais excede a capacidade "
                        "do template privado."
                    )
                _replace_scalar_markers(document, model)
                _populate_cover(document, model)
                _populate_summary_tables(document, model, table_slots)
                _populate_ergo(document, model, table_slots, capacity)
                _populate_psychosocial_images(
                    document,
                    model,
                    image_slots,
                    capacity,
                )
                _populate_technical_text(
                    document,
                    model,
                    manifest,
                    capacity,
                )
                _populate_question_tables(
                    document,
                    model,
                    table_slots,
                    capacity,
                )
                _populate_priorities_and_actions(
                    document,
                    model,
                    table_slots,
                )
                _assert_no_template_markers(document)
            else:
                document = _build_fallback_document(model)
            _set_document_metadata(document, model)
            document.save(output)
        except Exception as exc:
            if output.exists():
                output.unlink()
            raise DocumentAssemblyError(
                "Não foi possível montar o documento Word editável."
            ) from exc
        return output


def assemble_document(
    model: NormalizedAEP,
    output_path: str | Path,
    *,
    template_path: str | Path | None = None,
    manifest_path: str | Path | None = None,
) -> Path:
    return DocumentAssembler(
        template_path,
        manifest_path,
    ).assemble(model, output_path)
