from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


class RenderError(RuntimeError):
    pass


class LegacyConversionError(RuntimeError):
    """A binary Word document could not be converted in the isolated profile."""


@dataclass(slots=True)
class RenderResult:
    renderer: str
    pdf_path: Path | None
    page_images: list[Path] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def find_libreoffice(explicit: Path | None = None) -> Path | None:
    candidates = [
        explicit,
        Path(os.environ["AEP_LIBREOFFICE_PATH"])
        if os.getenv("AEP_LIBREOFFICE_PATH")
        else None,
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
    ]
    discovered = shutil.which("soffice") or shutil.which("libreoffice")
    if discovered:
        candidates.append(Path(discovered))
    for candidate in candidates:
        if candidate and candidate.is_file():
            return candidate.resolve()
    return None


def find_pdftoppm(explicit: Path | None = None) -> Path | None:
    if explicit and explicit.is_file():
        return explicit.resolve()
    discovered = shutil.which("pdftoppm")
    return Path(discovered).resolve() if discovered else None


def convert_legacy_doc_to_docx(
    input_path: Path,
    output_dir: Path,
    *,
    libreoffice_path: Path | None = None,
    timeout_seconds: int = 120,
) -> Path:
    """Convert a real OLE ``.doc`` with an isolated LibreOffice profile.

    The caller supplies server-owned paths.  LibreOffice receives no shell
    command, uses a fresh profile, and writes only inside ``output_dir``.
    """

    source = input_path.resolve()
    destination_dir = output_dir.resolve()
    if not source.is_file():
        raise LegacyConversionError("O documento binário não foi encontrado.")
    destination_dir.mkdir(parents=True, exist_ok=True)
    soffice = find_libreoffice(libreoffice_path)
    if soffice is None:
        raise LegacyConversionError(
            "O relatório Ergo é um DOC binário e o LibreOffice não está disponível."
        )

    with tempfile.TemporaryDirectory(
        prefix="aep-lo-doc-",
        dir=destination_dir,
        ignore_cleanup_errors=True,
    ) as profile_dir:
        profile_path = Path(profile_dir).resolve()
        profile_user = profile_path / "user"
        profile_user.mkdir()
        (profile_user / "registrymodifications.xcu").write_text(
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<oor:items xmlns:oor="http://openoffice.org/2001/registry">\n'
            '  <item oor:path="/org.openoffice.Office.Common/Security/Scripting">\n'
            '    <prop oor:name="MacroSecurityLevel" oor:op="fuse">'
            "<value>3</value></prop>\n"
            '    <prop oor:name="DisableMacrosExecution" oor:op="fuse">'
            "<value>true</value></prop>\n"
            "  </item>\n"
            "</oor:items>\n",
            encoding="utf-8",
        )
        profile_uri = profile_path.as_uri()
        command = [
            str(soffice),
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--norestore",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            'docx:Office Open XML Text',
            "--outdir",
            str(destination_dir),
            str(source),
        ]
        try:
            process = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=timeout_seconds,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise LegacyConversionError(
                "O LibreOffice não conseguiu converter o DOC binário com segurança."
            ) from exc

    converted = destination_dir / f"{source.stem}.docx"
    if (
        process.returncode != 0
        or not converted.is_file()
        or converted.stat().st_size == 0
        or not zipfile.is_zipfile(converted)
    ):
        if converted.is_file():
            converted.unlink()
        raise LegacyConversionError(
            "O LibreOffice não conseguiu converter o DOC binário com segurança."
        )
    try:
        with zipfile.ZipFile(converted) as archive:
            if "word/document.xml" not in archive.namelist():
                raise LegacyConversionError(
                    "A conversão do DOC binário não produziu um DOCX válido."
                )
    except zipfile.BadZipFile as exc:
        converted.unlink(missing_ok=True)
        raise LegacyConversionError(
            "A conversão do DOC binário não produziu um DOCX válido."
        ) from exc
    return converted


def inspect_docx_structure(path: Path) -> dict[str, int]:
    if not path.is_file() or not zipfile.is_zipfile(path):
        raise RenderError("O documento gerado não é um pacote DOCX válido.")
    with zipfile.ZipFile(path) as archive:
        if "word/document.xml" not in archive.namelist():
            raise RenderError("O pacote não contém word/document.xml.")
        media = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
    document = Document(path)
    return {
        "sections": len(document.sections),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "media_parts": len(media),
    }


def render_docx(
    input_path: Path,
    output_dir: Path,
    *,
    libreoffice_path: Path | None = None,
    word_script: Path | None = None,
    pdftoppm_path: Path | None = None,
    allow_word_fallback: bool = True,
    timeout_seconds: int = 180,
) -> RenderResult:
    output_dir = output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    source = input_path.resolve()
    soffice = find_libreoffice(libreoffice_path)
    warnings: list[str] = []

    if soffice:
        pdf_path = _render_with_libreoffice(
            soffice, source, output_dir, timeout_seconds
        )
        renderer = "libreoffice"
    elif allow_word_fallback and os.name == "nt":
        script = word_script or (
            Path(__file__).resolve().parents[2] / "scripts" / "render_with_word.ps1"
        )
        if not script.is_file():
            raise RenderError("LibreOffice ausente e script de fallback do Word indisponível.")
        pdf_path = output_dir / f"{source.stem}.pdf"
        _render_with_word(script, source, pdf_path, timeout_seconds)
        renderer = "word"
        warnings.append(
            "LibreOffice não estava disponível; a renderização visual usou o Word local."
        )
    else:
        raise RenderError("Nenhum renderizador DOCX está disponível.")

    rasterizer = find_pdftoppm(pdftoppm_path)
    pages: list[Path] = []
    if rasterizer:
        prefix = output_dir / "page"
        command = [
            str(rasterizer),
            "-png",
            "-r",
            "144",
            str(pdf_path),
            str(prefix),
        ]
        process = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            check=False,
        )
        if process.returncode != 0:
            raise RenderError("Falha ao converter o PDF em imagens.")
        pages = sorted(output_dir.glob("page-*.png"))
    else:
        warnings.append(
            "pdftoppm não foi localizado; o PDF foi criado, mas as páginas PNG não."
        )

    return RenderResult(
        renderer=renderer,
        pdf_path=pdf_path,
        page_images=pages,
        warnings=warnings,
    )


def _render_with_libreoffice(
    soffice: Path,
    source: Path,
    output_dir: Path,
    timeout_seconds: int,
) -> Path:
    with tempfile.TemporaryDirectory(
        prefix="aep-lo-",
        dir=output_dir,
        # LibreOffice can release extension-registry handles a fraction after
        # the conversion process exits on Windows. The rendered PDF is already
        # complete at that point, so a delayed profile cleanup must not turn a
        # successful render into an application error.
        ignore_cleanup_errors=True,
    ) as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        command = [
            str(soffice),
            "--headless",
            "--nologo",
            "--nodefault",
            "--nolockcheck",
            "--norestore",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(source),
        ]
        process = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=timeout_seconds,
            check=False,
        )
    pdf_path = output_dir / f"{source.stem}.pdf"
    if process.returncode != 0 or not pdf_path.is_file() or pdf_path.stat().st_size == 0:
        raise RenderError("O LibreOffice não conseguiu renderizar o documento.")
    return pdf_path


def _render_with_word(
    script: Path,
    source: Path,
    pdf_path: Path,
    timeout_seconds: int,
) -> None:
    process = subprocess.run(
        [
            "powershell.exe",
            "-NoProfile",
            "-STA",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(script),
            "-InputDocx",
            str(source),
            "-OutputPdf",
            str(pdf_path),
        ],
        capture_output=True,
        text=True,
        timeout=timeout_seconds,
        check=False,
    )
    if process.returncode != 0 or not pdf_path.is_file() or pdf_path.stat().st_size == 0:
        raise RenderError("O Word não conseguiu renderizar o documento.")
