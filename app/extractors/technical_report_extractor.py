"""Approved technical-report extraction without prose generation or rewriting."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
import re
from typing import Iterator

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models import (
    ActionPlanItem,
    AnalysisMode,
    PriorityItem,
    TechnicalAnalysis,
    TechnicalCategory,
    TechnicalReport,
    TechnicalSection,
)
from app.services.normalization import (
    classify_technical_heading,
    clean_text,
    extract_percentage,
    looks_like_heading,
    normalize_key,
    parse_ghe_reference,
)

from . import ExtractionError, UnsupportedSourceError


@dataclass(slots=True)
class _Unit:
    kind: str
    text: str = ""
    rows: list[list[str]] = field(default_factory=list)
    heading: bool = False


@dataclass(slots=True)
class _SectionBuilder:
    title: str
    source_role: str
    ghe_code: str | None
    ghe_name: str | None
    category: TechnicalCategory
    paragraphs: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)


def _iter_blocks(document: _Document) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _read_units(source: Path) -> list[_Unit]:
    try:
        if source.read_bytes()[:4] != b"PK\x03\x04":
            raise UnsupportedSourceError(
                "O relatório técnico deve ser um DOCX válido."
            )
        document = Document(source)
    except UnsupportedSourceError:
        raise
    except Exception as exc:
        raise ExtractionError("O relatório técnico DOCX está corrompido.") from exc

    units: list[_Unit] = []
    for block in _iter_blocks(document):
        if isinstance(block, Paragraph):
            text = clean_text(block.text)
            if not text:
                continue
            style_name = block.style.name if block.style else None
            style_key = normalize_key(style_name)
            text_key = normalize_key(text)
            code, _ = parse_ghe_reference(text)
            known_heading = (
                text_key
                in {
                    "visao geral",
                    "pontos positivos",
                    "pontos criticos",
                    "indicacoes de melhoria",
                }
                or "perguntas de maior relevancia" in text_key
            )
            numbered_heading = bool(
                re.match(r"^\d+\.\s+\S", text)
                and len(text) <= 140
            )
            ghe_heading = bool(
                code
                and len(text) <= 100
                and not style_key.startswith(("list", "lista"))
            )
            units.append(
                _Unit(
                    kind="paragraph",
                    text=text,
                    heading=bool(
                        known_heading
                        or numbered_heading
                        or ghe_heading
                        or style_key.startswith(
                            ("heading", "titulo", "title", "cabecalho")
                        )
                    ),
                )
            )
        else:
            rows = [
                [clean_text(cell.text) for cell in row.cells]
                for row in block.rows
            ]
            rows = [row for row in rows if any(row)]
            if rows:
                units.append(_Unit(kind="table", rows=rows))
    return units


def _category(title: str) -> TechnicalCategory:
    return TechnicalCategory(classify_technical_heading(title))


def _segment(
    units: list[_Unit], source_role: str, starting_order: int
) -> list[TechnicalSection]:
    builders: list[_SectionBuilder] = []
    current: _SectionBuilder | None = None
    active_ghe_code: str | None = None
    active_ghe_name: str | None = None

    for unit in units:
        if unit.kind == "paragraph" and unit.heading:
            code, name = parse_ghe_reference(unit.text)
            if code:
                active_ghe_code = code
                active_ghe_name = name
            elif re.match(r"^\d+\.\s+\S", unit.text):
                active_ghe_code = None
                active_ghe_name = None
            current = _SectionBuilder(
                title=unit.text,
                source_role=source_role,
                ghe_code=active_ghe_code,
                ghe_name=active_ghe_name,
                category=_category(unit.text),
            )
            builders.append(current)
            continue
        if current is None:
            current = _SectionBuilder(
                title="Conteúdo inicial",
                source_role=source_role,
                ghe_code=active_ghe_code,
                ghe_name=active_ghe_name,
                category=TechnicalCategory.OTHER,
            )
            builders.append(current)
        if unit.kind == "paragraph":
            current.paragraphs.append(unit.text)
        else:
            current.tables.append(unit.rows)

    sections: list[TechnicalSection] = []
    for local_order, builder in enumerate(builders):
        order = starting_order + local_order
        sections.append(
            TechnicalSection(
                section_id=f"technical-section-{order:04d}",
                order=order,
                title=builder.title,
                category=builder.category,
                source_role=builder.source_role,
                ghe_code_hint=builder.ghe_code,
                ghe_name_hint=builder.ghe_name,
                paragraphs=builder.paragraphs,
                tables=builder.tables,
            )
        )
    return sections


def _section_values(section: TechnicalSection) -> Iterator[str]:
    yield section.title
    yield from section.paragraphs
    for table in section.tables:
        for row in table:
            yield from (cell for cell in row if cell)


def _after_colon(value: str) -> str | None:
    for separator in (":", "–", "—"):
        if separator in value:
            candidate = clean_text(value.split(separator, 1)[1])
            if candidate:
                return candidate
    return None


def _build_analyses(
    sections: list[TechnicalSection],
) -> list[TechnicalAnalysis]:
    keys: list[str] = []
    grouped: dict[str, list[TechnicalSection]] = {}
    names: dict[str, str | None] = {}
    for section in sections:
        if not section.ghe_code_hint:
            continue
        key = section.ghe_code_hint
        if key not in grouped:
            grouped[key] = []
            keys.append(key)
            names[key] = section.ghe_name_hint
        elif not names[key] and section.ghe_name_hint:
            names[key] = section.ghe_name_hint
        grouped[key].append(section)

    analyses: list[TechnicalAnalysis] = []
    for order, key in enumerate(keys):
        group_sections = grouped[key]
        favorable: str | None = None
        classification: str | None = None
        readings: list[str] = []
        for section in group_sections:
            if section.category == TechnicalCategory.FAVORABILITY and not favorable:
                favorable = next(
                    (
                        percentage
                        for value in _section_values(section)
                        if (percentage := extract_percentage(value))
                    ),
                    None,
                )
            if (
                section.category == TechnicalCategory.CLASSIFICATION
                and not classification
            ):
                classification = _after_colon(section.title)
                if not classification:
                    classification = next(
                        (
                            value
                            for value in section.paragraphs
                            if clean_text(value)
                        ),
                        None,
                    )
            if section.category == TechnicalCategory.TECHNICAL_READING:
                readings.extend(section.paragraphs)
        analyses.append(
            TechnicalAnalysis(
                analysis_id=f"technical-analysis-{order:03d}",
                order=order,
                ghe_code_hint=key,
                ghe_name_hint=names[key],
                sections=group_sections,
                favorable_percentage=favorable,
                classification=classification,
                technical_reading=readings,
            )
        )
    return analyses


def _priorities(sections: list[TechnicalSection]) -> list[PriorityItem]:
    result: list[PriorityItem] = []
    for section in sections:
        if section.category != TechnicalCategory.PRIORITIZATION:
            continue
        for table in section.tables:
            if not table:
                continue
            headers = {
                normalize_key(cell): index
                for index, cell in enumerate(table[0])
            }
            order_index = next(
                (index for key, index in headers.items() if key == "ordem"),
                None,
            )
            ghe_index = next(
                (index for key, index in headers.items() if key == "ghe"),
                None,
            )
            class_index = next(
                (
                    index
                    for key, index in headers.items()
                    if "classificacao" in key or "prioridade" in key
                ),
                None,
            )
            reason_index = next(
                (
                    index
                    for key, index in headers.items()
                    if "justificativa" in key
                    or "fundamentacao" in key
                    or "descricao" in key
                ),
                None,
            )
            if ghe_index is None or reason_index is None:
                continue
            for row in table[1:]:
                if not any(row):
                    continue
                ghe_value = clean_text(row[ghe_index])
                reason = clean_text(row[reason_index])
                if not ghe_value or not reason:
                    continue
                result.append(
                    PriorityItem(
                        order=len(result),
                        text=reason,
                        ghe_code_hint=ghe_value,
                        level=(
                            clean_text(row[class_index])
                            if class_index is not None
                            else None
                        ),
                        source_section_id=section.section_id,
                    )
                )
    return result


def _header_indexes(row: list[str]) -> dict[str, int]:
    result: dict[str, int] = {}
    for index, cell in enumerate(row):
        key = normalize_key(cell)
        if "evolucao" in key or "registro" in key:
            result.setdefault("evolution", index)
        elif "responsavel" in key:
            result.setdefault("responsible", index)
        elif "prazo" in key or "data" == key:
            result.setdefault("deadline", index)
        elif "indicador" in key:
            result.setdefault("indicator", index)
        elif "prioridade" in key:
            result.setdefault("priority", index)
        elif key == "ghe" or "ghe " in key or "grupo" in key:
            result.setdefault("ghe", index)
        elif "acao" in key or "medida" in key or "atividade" in key:
            result.setdefault("action", index)
    return result


def _at(row: list[str], indexes: dict[str, int], key: str) -> str | None:
    index = indexes.get(key)
    if index is None or index >= len(row):
        return None
    return clean_text(row[index]) or None


def _action_plan(sections: list[TechnicalSection]) -> list[ActionPlanItem]:
    result: list[ActionPlanItem] = []
    for section in sections:
        if section.category != TechnicalCategory.ACTION_PLAN:
            continue
        for table in section.tables:
            if not table:
                continue
            indexes = _header_indexes(table[0])
            has_header = "action" in indexes
            data_rows = table[1:] if has_header else table
            for row in data_rows:
                nonempty = [cell for cell in row if cell]
                if not nonempty:
                    continue
                action = (
                    _at(row, indexes, "action") if has_header else nonempty[0]
                )
                if not action:
                    continue
                ghe_hint = _at(row, indexes, "ghe") or section.ghe_code_hint
                ghe_code, _ = parse_ghe_reference(ghe_hint)
                result.append(
                    ActionPlanItem(
                        order=len(result),
                        action=action,
                        ghe_code_hint=ghe_hint,
                        responsible=_at(row, indexes, "responsible"),
                        deadline=_at(row, indexes, "deadline"),
                        indicator=_at(row, indexes, "indicator"),
                        priority=_at(row, indexes, "priority"),
                        source_section_id=section.section_id,
                        evolution_records="",
                    )
                )
    return result


def _conclusion(sections: list[TechnicalSection]) -> list[str]:
    paragraphs: list[str] = []
    for section in sections:
        if section.category != TechnicalCategory.CONCLUSION:
            continue
        paragraphs.extend(section.paragraphs)
    return paragraphs


class TechnicalReportExtractor:
    """Extract either one integrated report or the two approved agent reports."""

    def extract(
        self,
        *,
        integrated_path: str | Path | None = None,
        psychosocial_path: str | Path | None = None,
        ergonomic_path: str | Path | None = None,
    ) -> TechnicalReport:
        integrated = Path(integrated_path) if integrated_path is not None else None
        psychosocial = (
            Path(psychosocial_path) if psychosocial_path is not None else None
        )
        ergonomic = Path(ergonomic_path) if ergonomic_path is not None else None
        if integrated and (psychosocial or ergonomic):
            raise ExtractionError(
                "Escolha o relatório integrado ou os dois relatórios separados."
            )
        if integrated:
            sources = [("integrated", integrated)]
            mode = AnalysisMode.INTEGRATED
        elif psychosocial and ergonomic:
            sources = [
                ("psychosocial_agent", psychosocial),
                ("ergonomic_agent", ergonomic),
            ]
            mode = AnalysisMode.SEPARATE
        else:
            raise ExtractionError(
                "O modo separado requer os relatórios psicossocial e ergonômico."
            )

        sections: list[TechnicalSection] = []
        for source_role, path in sources:
            units = _read_units(path)
            sections.extend(_segment(units, source_role, len(sections)))

        warnings: list[str] = []
        categories = {section.category for section in sections}
        for category, label in (
            (TechnicalCategory.PRIORITIZATION, "priorização"),
            (TechnicalCategory.ACTION_PLAN, "plano de ação"),
            (TechnicalCategory.CONCLUSION, "conclusão"),
        ):
            if category not in categories:
                warnings.append(
                    f"O relatório técnico não contém seção reconhecida de {label}."
                )

        return TechnicalReport(
            mode=mode,
            sections=sections,
            analyses=_build_analyses(sections),
            priorities=_priorities(sections),
            action_plan=_action_plan(sections),
            conclusion=_conclusion(sections),
            source_roles=[source_role for source_role, _ in sources],
            warnings=warnings,
        )


def extract_technical_report(
    *,
    integrated_path: str | Path | None = None,
    psychosocial_path: str | Path | None = None,
    ergonomic_path: str | Path | None = None,
) -> TechnicalReport:
    return TechnicalReportExtractor().extract(
        integrated_path=integrated_path,
        psychosocial_path=psychosocial_path,
        ergonomic_path=ergonomic_path,
    )
