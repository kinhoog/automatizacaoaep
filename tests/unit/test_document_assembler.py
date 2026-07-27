from __future__ import annotations

import hashlib
import json
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn

from app.services.document_assembler import (
    DocumentAssembler,
    DocumentAssemblyError,
)
from app.services.normalization import normalize_key


def _all_text(document: Document) -> str:
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *cells])


def test_generates_editable_docx_without_optional_logo(
    normalized_model,
    tmp_path: Path,
) -> None:
    output = tmp_path / "Documento AEP - SINTETICO.docx"
    result = DocumentAssembler().assemble(normalized_model, output)

    assert result == output
    assert output.is_file()
    assert output.read_bytes().startswith(b"PK\x03\x04")

    document = Document(output)
    text = _all_text(document)
    assert normalized_model.company.legal_name in text
    assert "Hierarquia dos GHEs" in text
    assert "GHE 10" in text
    assert "Resultados do Ergo" in text
    assert "Conclusao sintetica" in text

    # Card + hierarchy image. A third image would indicate an optional logo.
    assert len(document.inline_shapes) == 2
    assert document.settings.element.find(qn("w:updateFields")) is not None


def test_generated_action_log_column_remains_blank(
    normalized_model,
    tmp_path: Path,
) -> None:
    output = tmp_path / "acao-editavel.docx"
    DocumentAssembler().assemble(normalized_model, output)
    document = Document(output)
    action_table = next(
        table
        for table in document.tables
        if table.rows
        and normalize_key(table.rows[0].cells[-1].text)
        == "evolucao registros"
    )

    assert len(action_table.rows) > 1
    assert all(
        not row.cells[-1].text.strip() for row in action_table.rows[1:]
    )


def test_generated_docx_changes_when_input_changes(
    normalized_model,
    tmp_path: Path,
) -> None:
    first = tmp_path / "primeiro.docx"
    second = tmp_path / "segundo.docx"
    DocumentAssembler().assemble(normalized_model, first)

    mutated = normalized_model.model_copy(deep=True)
    mutated.company.legal_name = "EMPRESA SINTETICA MUTACAO LTDA"
    mutated.official_ghes[0].population += 2
    DocumentAssembler().assemble(mutated, second)

    first_document = Document(first)
    second_document = Document(second)
    assert _all_text(first_document) != _all_text(second_document)
    assert "EMPRESA SINTETICA MUTACAO LTDA" in _all_text(second_document)
    assert first.read_bytes() != second.read_bytes()


def test_private_template_rejects_manifest_without_sanitization_proof(
    normalized_model,
    public_fixtures: Path,
    tmp_path: Path,
) -> None:
    template = public_fixtures / "template_aep_sintetico.docx"
    structure = Document(template)
    manifest = tmp_path / "template.manifest.json"
    payload = {
        "schema_version": 1,
        "template_sha256": hashlib.sha256(template.read_bytes()).hexdigest(),
        "paragraph_count": len(structure.paragraphs),
        "table_count": len(structure.tables),
        "inline_shape_count": len(structure.inline_shapes),
        "slots": {
            "tables": {},
            "body_image_order": {},
        },
    }
    manifest.write_text(
        json.dumps(payload, ensure_ascii=False),
        encoding="utf-8",
    )

    rejected = tmp_path / "manifesto-sem-prova.docx"
    with pytest.raises(DocumentAssemblyError):
        DocumentAssembler(template, manifest).assemble(
            normalized_model,
            rejected,
        )
    assert not rejected.exists()


def test_configured_private_template_requires_manifest(
    normalized_model,
    public_fixtures: Path,
    tmp_path: Path,
) -> None:
    output = tmp_path / "sem-manifesto.docx"

    with pytest.raises(DocumentAssemblyError):
        DocumentAssembler(
            public_fixtures / "template_aep_sintetico.docx"
        ).assemble(normalized_model, output)

    assert not output.exists()


def test_missing_configured_template_never_uses_fallback(
    normalized_model,
    tmp_path: Path,
) -> None:
    output = tmp_path / "fallback-indevido.docx"

    with pytest.raises(DocumentAssemblyError):
        DocumentAssembler(
            tmp_path / "template-inexistente.docx",
            tmp_path / "manifesto-inexistente.json",
        ).assemble(normalized_model, output)

    assert not output.exists()
