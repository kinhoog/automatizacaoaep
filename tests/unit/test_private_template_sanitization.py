from __future__ import annotations

import hashlib
import io
import json
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches
from PIL import Image

from app.services.document_assembler import (
    DocumentAssembler,
    DocumentAssemblyError,
)
from app.services.validation import validate_normalized_aep
from scripts.prepare_private_template import prepare


_RESIDUE = "RESIDUO_SINTETICO_PRIVADO_SEM_VALIDADE"


def _all_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join([*paragraphs, *cells])


def _png(index: int) -> bytes:
    image = Image.new(
        "RGB",
        (24 + index, 20 + index),
        ((index * 31) % 255, (index * 67) % 255, (index * 97) % 255),
    )
    stream = io.BytesIO()
    image.save(stream, format="PNG")
    return stream.getvalue()


def _fill_table(table, *, prefix: str) -> None:
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.text = f"{prefix}_{row_index}_{column_index}"


def _build_synthetic_reference(path: Path) -> None:
    document = Document()
    document.core_properties.title = f"{_RESIDUE}_METADADO"
    document.add_paragraph("AVALIAÇÃO ERGONÔMICA PRELIMINAR")
    document.add_paragraph("EMPRESA SINTÉTICA DE ORIGEM — SEM VALIDADE")
    document.add_paragraph("JANEIRO/2000 — SEM VALIDADE")
    document.add_paragraph("Identificação da empresa")

    # The retained pilot layout has eleven body image positions. Each source
    # image is unique so the test can prove that dynamic media was removed.
    for index in range(11):
        paragraph = document.add_paragraph()
        paragraph.add_run().add_picture(
            io.BytesIO(_png(index)),
            width=Inches(0.25),
        )

    revision = document.add_table(rows=4, cols=3)
    _fill_table(revision, prefix="HISTORICO")
    for row in revision.rows[2:]:
        for cell in row.cells:
            cell.text = _RESIDUE

    hierarchy = document.add_table(rows=5, cols=2)
    hierarchy.rows[0].cells[0].text = "GHE"
    hierarchy.rows[0].cells[1].text = "População"
    for index in range(1, 4):
        hierarchy.rows[index].cells[0].text = (
            f"GHE {index:02d} - {_RESIDUE}_{index}"
        )
        hierarchy.rows[index].cells[1].text = str(index)
    hierarchy.rows[4].cells[0].text = "TOTAL"
    hierarchy.rows[4].cells[1].text = _RESIDUE

    summary = document.add_table(rows=4, cols=2)
    for index, row in enumerate(summary.rows):
        row.cells[0].text = f"Resumo {index + 1}"
        row.cells[1].text = _RESIDUE

    for slot in range(3):
        title = document.add_table(rows=1, cols=2)
        title.rows[0].cells[0].text = "eProtege ERGO"
        title.rows[0].cells[1].text = ""

        metadata = document.add_table(rows=1, cols=4)
        metadata.rows[0].cells[0].text = "Empresa"
        metadata.rows[0].cells[1].text = _RESIDUE
        metadata.rows[0].cells[2].text = "Setor / GHE"
        metadata.rows[0].cells[3].text = f"{_RESIDUE}_{slot + 1}"

        counts = document.add_table(rows=1, cols=3)
        counts.rows[0].cells[0].text = _RESIDUE
        counts.rows[0].cells[1].text = "Indicadores"
        counts.rows[0].cells[2].text = _RESIDUE

        questions = document.add_table(rows=3, cols=3)
        questions.rows[0].cells[0].text = "Pergunta"
        questions.rows[0].cells[1].text = "Resposta"
        questions.rows[0].cells[2].text = "Observação / Orientação"
        for row in questions.rows[1:]:
            for cell in row.cells:
                cell.text = _RESIDUE

    for slot in range(3):
        questions = document.add_table(rows=3, cols=4)
        questions.rows[0].cells[0].text = "Pergunta"
        questions.rows[0].cells[1].text = "Favorável"
        questions.rows[0].cells[2].text = "Classe"
        questions.rows[0].cells[3].text = "Leitura técnica"
        for row in questions.rows[1:]:
            for cell in row.cells:
                cell.text = f"{_RESIDUE}_TECNICO_{slot + 1}"

    priorities = document.add_table(rows=3, cols=4)
    priorities.rows[0].cells[0].text = "Ordem"
    priorities.rows[0].cells[1].text = "GHE"
    priorities.rows[0].cells[2].text = "Nível"
    priorities.rows[0].cells[3].text = "Prioridade"
    for row in priorities.rows[1:]:
        for cell in row.cells:
            cell.text = _RESIDUE

    actions = document.add_table(rows=3, cols=5)
    actions.rows[0].cells[0].text = "Prioridade"
    actions.rows[0].cells[1].text = "GHE/Abrangência"
    actions.rows[0].cells[2].text = "Plano de ação"
    actions.rows[0].cells[3].text = "Responsável"
    actions.rows[0].cells[4].text = "Evolução/Registros"
    for row in actions.rows[1:]:
        for cell in row.cells:
            cell.text = _RESIDUE

    document.add_paragraph("Resultado dos Riscos Psicossociais por GHE")
    for slot in range(1, 4):
        document.add_paragraph(f"GHE {slot:02d} — {_RESIDUE}_{slot}")
        for heading in (
            "Visão geral",
            "Pontos positivos",
            "Pontos críticos",
            "Indicações de melhoria",
        ):
            document.add_paragraph(heading)
            document.add_paragraph(f"{_RESIDUE}_{slot}_{heading}")
        document.add_paragraph("Perguntas de maior relevância no GHE")
    document.add_paragraph("Priorizações recomendadas")
    document.add_paragraph("Plano de ação geral integrado")
    document.add_paragraph("Conclusão Técnica")
    document.add_paragraph(f"{_RESIDUE}_CONCLUSAO")
    document.add_paragraph("Termo de encerramento")
    document.save(path)


@pytest.fixture
def sanitized_template(tmp_path: Path) -> tuple[Path, Path, dict]:
    source = tmp_path / "referencia_sintetica_privada.docx"
    template = tmp_path / "template_sintetico_sanitizado.docx"
    _build_synthetic_reference(source)
    source_hash = hashlib.sha256(source.read_bytes()).hexdigest()

    prepared, manifest_path = prepare(source, template)

    assert prepared == template
    assert hashlib.sha256(source.read_bytes()).hexdigest() == source_hash
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    return template, manifest_path, manifest


def test_template_preparation_proves_text_and_media_sanitization(
    sanitized_template: tuple[Path, Path, dict],
) -> None:
    template, _, manifest = sanitized_template
    sanitization = manifest["sanitization"]

    assert sanitization["status"] == "sanitized"
    assert sanitization["contract_version"] == 1
    assert sanitization["marker_count"] > 0
    assert sanitization["capacity"] == {
        "official_ghes": 3,
        "ergo_blocks": 3,
        "psychosocial_ghe_pairs": 3,
        "technical_ghe_sections": 3,
    }
    assert _RESIDUE not in _all_text(template)

    with zipfile.ZipFile(template) as package:
        media_hashes = {
            hashlib.sha256(package.read(name)).hexdigest()
            for name in package.namelist()
            if name.startswith("word/media/")
        }
    assert not (
        set(sanitization["source_dynamic_media_sha256"]) & media_hashes
    )


def test_retained_template_assembly_has_no_marker_or_source_residue(
    normalized_model,
    sanitized_template: tuple[Path, Path, dict],
    tmp_path: Path,
) -> None:
    template, manifest, _ = sanitized_template
    output = tmp_path / "aep_sanitizado.docx"

    DocumentAssembler(template, manifest).assemble(normalized_model, output)

    text = _all_text(output)
    assert normalized_model.company.legal_name in text
    assert "{{AEP_" not in text
    assert _RESIDUE not in text


def test_retained_template_clears_unused_and_partially_extracted_slots(
    normalized_model,
    sanitized_template: tuple[Path, Path, dict],
    tmp_path: Path,
) -> None:
    template, manifest, _ = sanitized_template
    partial = normalized_model.model_copy(deep=True)
    partial.official_ghes = partial.official_ghes[:2]
    kept_codes = {ghe.canonical_code for ghe in partial.official_ghes}
    partial.technical.analyses = [
        analysis
        for analysis in partial.technical.analyses
        if (analysis.official_ghe_code or analysis.ghe_code_hint) in kept_codes
    ]
    partial.psychosocial.images = [
        image
        for image in partial.psychosocial.images
        if image.official_ghe_code is None
        or image.official_ghe_code in kept_codes
    ]
    partial.psychosocial.blocks = [
        block
        for block in partial.psychosocial.blocks
        if block.official_ghe_code is None
        or block.official_ghe_code in kept_codes
    ]
    kept_items = [
        item
        for item in partial.reconciliation.items
        if item.official_ghe_code in kept_codes
    ]
    kept_source_ids = {item.source_id for item in kept_items}
    partial.reconciliation.items = kept_items
    partial.ergo.blocks = [
        block
        for block in partial.ergo.blocks
        if block.source_id in kept_source_ids
    ]

    output = tmp_path / "aep_dois_ghes.docx"
    DocumentAssembler(template, manifest).assemble(partial, output)

    text = _all_text(output)
    assert "{{AEP_" not in text
    assert _RESIDUE not in text
    assert partial.official_ghes[0].canonical_code in text
    assert partial.official_ghes[1].canonical_code in text


def test_retained_template_supports_different_codes_without_legacy_slots(
    normalized_model,
    sanitized_template: tuple[Path, Path, dict],
    tmp_path: Path,
) -> None:
    template, manifest, _ = sanitized_template
    changed = normalized_model.model_copy(deep=True)
    mapping = {
        ghe.canonical_code: f"GHE {index:02d}"
        for index, ghe in enumerate(changed.official_ghes, start=41)
    }
    for ghe in changed.official_ghes:
        old_code = ghe.canonical_code
        ghe.code = mapping[old_code]
        ghe.name = f"GRUPO SINTÉTICO ALTERADO {mapping[old_code][-2:]}"
    for analysis in changed.technical.analyses:
        code = analysis.official_ghe_code or analysis.ghe_code_hint
        if code in mapping:
            analysis.official_ghe_code = mapping[code]
            analysis.ghe_code_hint = mapping[code]
    for image in changed.psychosocial.images:
        if image.official_ghe_code in mapping:
            image.official_ghe_code = mapping[image.official_ghe_code]
    for block in changed.psychosocial.blocks:
        if block.official_ghe_code in mapping:
            block.official_ghe_code = mapping[block.official_ghe_code]
    for item in changed.reconciliation.items:
        if item.official_ghe_code in mapping:
            item.official_ghe_code = mapping[item.official_ghe_code]

    output = tmp_path / "aep_codigos_alterados.docx"
    DocumentAssembler(template, manifest).assemble(changed, output)

    text = _all_text(output)
    assert "{{AEP_" not in text
    assert _RESIDUE not in text
    assert all(code in text for code in mapping.values())


def test_retained_template_never_truncates_more_ghes_than_capacity(
    normalized_model,
    sanitized_template: tuple[Path, Path, dict],
    tmp_path: Path,
) -> None:
    template, manifest, _ = sanitized_template
    too_many = normalized_model.model_copy(deep=True)
    fourth = too_many.official_ghes[0].model_copy(deep=True)
    fourth.code = "GHE 99"
    fourth.name = "GRUPO SINTÉTICO EXCEDENTE"
    too_many.official_ghes.append(fourth)
    output = tmp_path / "nao_deve_existir.docx"

    with pytest.raises(DocumentAssemblyError):
        DocumentAssembler(template, manifest).assemble(too_many, output)

    assert not output.exists()


def test_partial_technical_or_image_coverage_blocks_final_generation(
    normalized_model,
) -> None:
    partial = normalized_model.model_copy(deep=True)
    missing_code = partial.official_ghes[-1].canonical_code
    partial.technical.analyses = [
        analysis
        for analysis in partial.technical.analyses
        if (analysis.official_ghe_code or analysis.ghe_code_hint) != missing_code
    ]
    partial.psychosocial.images = [
        image
        for image in partial.psychosocial.images
        if image.official_ghe_code != missing_code
    ]

    preliminary = validate_normalized_aep(
        partial,
        require_complete_reconciliation=False,
    )
    final = validate_normalized_aep(
        partial,
        require_complete_reconciliation=True,
    )
    preliminary_codes = {issue.code for issue in preliminary.warnings}
    final_codes = {issue.code for issue in final.errors}

    assert "technical_ghe_analysis_missing" in preliminary_codes
    assert "psychosocial_ghe_image_missing" in preliminary_codes
    assert "technical_ghe_analysis_missing" in final_codes
    assert "psychosocial_ghe_image_missing" in final_codes
