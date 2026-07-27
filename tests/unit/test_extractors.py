from __future__ import annotations

from pathlib import Path
import re
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from openpyxl import Workbook, load_workbook
import pytest

from app.extractors.ghe_extractor import extract_ghes
from app.extractors.technical_report_extractor import (
    extract_technical_report,
)
from app.models import AnalysisMode, ContentKind, FileKind, ImageRole


def test_ghe_spreadsheet_is_source_of_truth_without_person_fields(
    ghe_result,
) -> None:
    assert [ghe.canonical_code for ghe in ghe_result.ghes] == [
        "GHE 10",
        "GHE 20",
        "GHE 30",
    ]
    assert [ghe.population for ghe in ghe_result.ghes] == [5, 8, 3]
    assert ghe_result.total_population == 16

    by_code = {ghe.canonical_code: ghe for ghe in ghe_result.ghes}
    assert by_code["GHE 10"].sectors == ["Unidade Alfa"]
    assert by_code["GHE 10"].roles == [
        "Operador de bancada",
        "Apoio de processo",
    ]
    assert by_code["GHE 20"].sectors == [
        "Unidade Beta",
        "Suporte remoto",
    ]
    assert by_code["GHE 30"].roles == ["Tecnico de laboratorio"]

    audit_payload = ghe_result.audit_dict()
    serialized_keys = {
        key.casefold()
        for ghe in audit_payload["ghes"]
        for key in ghe
    }
    assert not serialized_keys.intersection(
        {"person", "employee", "funcionario", "colaborador", "nome"}
    )
    assert all("source_rows" not in ghe for ghe in audit_payload["ghes"])


def test_ghe_spreadsheet_without_dimension_metadata_is_supported(
    tmp_path: Path,
) -> None:
    source = tmp_path / "ghe_sintetico.xlsx"
    dimensionless = tmp_path / "ghe_sintetico_sem_dimensao.xlsx"
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Resumo"
    summary.append(["ARQUIVO SINTETICO SEM VALIDADE", None])
    summary.append(["GHE", "QTD."])
    summary.append(["GHE 01 - Administrativo", 5])
    summary.append(["GHE 02 - Producao", 4])
    summary.append(["GHE 03 - Vendas", 3])
    for code, sector, role in (
        ("GHE 01", "Administrativo", "Assistente sintetico"),
        ("GHE 02", "Producao", "Operador sintetico"),
        ("GHE 03", "Vendas", "Vendedor sintetico"),
    ):
        detail = workbook.create_sheet(code)
        detail.append(["Setor", "Cargo"])
        detail.append([sector, role])
    workbook.save(source)

    with ZipFile(source) as original, ZipFile(
        dimensionless,
        "w",
        compression=ZIP_DEFLATED,
    ) as rewritten:
        for item in original.infolist():
            content = original.read(item.filename)
            if item.filename.startswith("xl/worksheets/sheet"):
                content = re.sub(
                    rb"<dimension\b[^>]*/>",
                    b"",
                    content,
                    count=1,
                )
            rewritten.writestr(item, content)

    read_only = load_workbook(dimensionless, read_only=True)
    try:
        assert read_only["Resumo"].max_row is None
    finally:
        read_only.close()

    result = extract_ghes(dimensionless)

    assert [ghe.canonical_code for ghe in result.ghes] == [
        "GHE 01",
        "GHE 02",
        "GHE 03",
    ]
    assert [ghe.population for ghe in result.ghes] == [5, 4, 3]
    assert result.total_population == 12
    assert result.ghes[1].sectors == ["Producao"]
    assert result.ghes[1].roles == ["Operador sintetico"]


def test_ergo_html_disguised_as_doc_preserves_blocks_and_visual_order(
    ergo_report,
) -> None:
    assert ergo_report.detected_format == FileKind.HTML_DOC
    assert [block.source_code for block in ergo_report.blocks] == [
        "GHE 10",
        "GHE 20",
        "GHE 30",
        "GHE 99",
    ]
    assert [block.order for block in ergo_report.blocks] == [0, 1, 2, 3]

    for block in ergo_report.blocks:
        element_orders = [element.order for element in block.elements]
        assert element_orders == sorted(element_orders)
        assert any(
            element.kind == ContentKind.TABLE for element in block.elements
        )
        assert any(
            element.kind == ContentKind.IMAGE for element in block.elements
        )
        assert block.questions
        assert block.answers
        assert block.observations
        assert block.guidance


def test_technical_docx_extracts_titles_tables_and_approved_text(
    integrated_report,
) -> None:
    titles = [section.title for section in integrated_report.sections]
    assert "Visao geral" in titles
    assert "Perguntas de maior relevancia" in titles
    assert "Plano de acao integrado" in titles
    assert "Conclusao tecnica" in titles

    tables = [
        table
        for section in integrated_report.sections
        for table in section.tables
    ]
    assert len(tables) >= 5
    assert any(
        "Pergunta" in " ".join(cell for row in table for cell in row)
        for table in tables
    )
    action_section = next(
        section
        for section in integrated_report.sections
        if section.title == "Plano de acao integrado"
    )
    assert action_section.tables
    assert action_section.tables[0][0][0] == "Acao"
    assert integrated_report.action_plan

    first_analysis = integrated_report.analyses[0]
    assert first_analysis.ghe_code_hint == "GHE 10"
    assert first_analysis.favorable_percentage == "72%"
    assert any(
        "Texto sintetico aprovado" in paragraph
        for section in first_analysis.sections
        for paragraph in section.paragraphs
    )


def test_technical_report_recognizes_normal_style_uppercase_sections(
    tmp_path: Path,
) -> None:
    source = tmp_path / "tecnico_sintetico_estilo_normal.docx"
    document = Document()
    document.add_paragraph("ARQUIVO SINTETICO SEM VALIDADE")
    document.add_paragraph(
        "ANALISE TECNICA POR GHE - GHE 01 - ADMINISTRATIVO"
    )
    document.add_paragraph("Visao geral")
    document.add_paragraph("Texto tecnico sintetico aprovado.")
    document.add_paragraph("PRIORIZACAO RECOMENDADA")
    priority = document.add_table(rows=2, cols=4)
    priority.rows[0].cells[0].text = "Ordem"
    priority.rows[0].cells[1].text = "GHE"
    priority.rows[0].cells[2].text = "Classificacao"
    priority.rows[0].cells[3].text = "Justificativa"
    priority.rows[1].cells[0].text = "1"
    priority.rows[1].cells[1].text = "GHE 01 - Administrativo"
    priority.rows[1].cells[2].text = "Prioridade moderada"
    priority.rows[1].cells[3].text = "Justificativa sintetica aprovada."
    document.add_paragraph(
        "PLANO DE ACAO GERAL INTEGRADO - ERGONOMIA E PSICOSSOCIAL"
    )
    action = document.add_table(rows=2, cols=5)
    action.rows[0].cells[0].text = "Prioridade"
    action.rows[0].cells[1].text = "GHE/Abrangencia"
    action.rows[0].cells[2].text = "Plano de acao"
    action.rows[0].cells[3].text = "Responsavel"
    action.rows[0].cells[4].text = "Evolucao/Registros"
    action.rows[1].cells[0].text = "Moderada"
    action.rows[1].cells[1].text = "GHE 01 - Administrativo"
    action.rows[1].cells[2].text = "Acao sintetica aprovada."
    action.rows[1].cells[3].text = "Responsavel sintetico"
    document.add_paragraph("CONCLUSAO TECNICA")
    document.add_paragraph("Texto final sintetico aprovado.")
    document.save(source)

    report = extract_technical_report(integrated_path=source)

    assert [item.ghe_code_hint for item in report.priorities] == [
        "GHE 01 - Administrativo"
    ]
    assert [item.action for item in report.action_plan] == [
        "Acao sintetica aprovada."
    ]
    assert report.conclusion == ["Texto final sintetico aprovado."]
    assert report.warnings == []


def test_psychosocial_images_are_associated_by_context_and_ghe(
    psychosocial_report,
) -> None:
    assert len(psychosocial_report.images) == 13
    assert [block.official_ghe_code for block in psychosocial_report.blocks] == [
        None,
        "GHE 10",
        "GHE 20",
        "GHE 30",
    ]
    assert len(psychosocial_report.blocks[0].images) == 1
    assert psychosocial_report.blocks[0].images[0].role == (
        ImageRole.GENERAL_PANEL
    )

    for block in psychosocial_report.blocks[1:]:
        assert len(block.images) == 4
        assert {image.role for image in block.images} == {
            ImageRole.GHE_PANEL,
            ImageRole.DOMAIN_SUMMARY,
            ImageRole.RADAR,
            ImageRole.RISK_MATRIX,
        }
        assert all(
            image.official_ghe_code == block.official_ghe_code
            for image in block.images
        )


def test_technical_report_supports_integrated_and_separate_modes(
    public_fixtures: Path,
    integrated_report,
) -> None:
    source = public_fixtures / "tecnico_integrado_sintetico.docx"
    separate = extract_technical_report(
        psychosocial_path=source,
        ergonomic_path=source,
    )

    assert integrated_report.mode == AnalysisMode.INTEGRATED
    assert integrated_report.source_roles == ["integrated"]
    assert separate.mode == AnalysisMode.SEPARATE
    assert separate.source_roles == [
        "psychosocial_agent",
        "ergonomic_agent",
    ]
    assert len(separate.sections) == 2 * len(integrated_report.sections)
    assert separate.conclusion == 2 * integrated_report.conclusion


def test_separate_mode_requires_both_agent_reports(
    public_fixtures: Path,
) -> None:
    source = public_fixtures / "tecnico_integrado_sintetico.docx"
    with pytest.raises(ValueError):
        extract_technical_report(psychosocial_path=source)
