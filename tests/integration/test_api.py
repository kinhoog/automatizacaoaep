from __future__ import annotations

import hashlib
import io
import json
import os
import shutil
import time
import zipfile
from dataclasses import replace
from pathlib import Path
from typing import Any

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import load_workbook
import pytest

import app.main as web_app
from app.services.pipeline import _compatibility_input_fingerprint


def _wait_for_terminal_status(
    client: TestClient, job_id: str, *, timeout: float = 15
) -> dict[str, Any]:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        response = client.get(f"/api/jobs/{job_id}")
        assert response.status_code == 200
        payload = response.json()
        if payload["status"] in {"completed", "failed"}:
            return payload
        time.sleep(0.02)
    raise AssertionError("A geração assíncrona não terminou no prazo do teste.")


def _extra_ergo_decision(validation: dict[str, Any]) -> dict[str, Any]:
    items = validation["reconciliation"]["items"]
    pending = next(
        item for item in items if item["status"] == "needs_review"
    )
    return {
        "source_id": pending["source_id"],
        "action": "not_applicable",
        "not_applicable": True,
        "reason": "Bloco sintético sem GHE oficial correspondente.",
    }


def _generate_and_download(
    client: TestClient,
    validation: dict[str, Any],
) -> bytes:
    job_id = validation["job_id"]
    response = client.post(
        "/api/generate",
        json={
            "job_id": job_id,
            "reconciliations": [_extra_ergo_decision(validation)],
        },
    )
    assert response.status_code == 202, response.text
    completed = _wait_for_terminal_status(client, job_id)
    assert completed["status"] == "completed", completed
    document_response = client.get(f"/api/jobs/{job_id}/document")
    assert document_response.status_code == 200
    return document_response.content


def _population_for_ghe(document_bytes: bytes, code: str) -> int:
    document = Document(io.BytesIO(document_bytes))
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if (
                values
                and values[0].startswith(code)
                and values[-1].isdigit()
            ):
                return int(values[-1])
    raise AssertionError(f"O GHE sintético {code} não foi encontrado na saída.")


def test_health_endpoint_reports_temporary_ready_pipeline(
    api_client: TestClient,
) -> None:
    response = api_client.get("/api/health")

    assert response.status_code == 200
    assert response.json() == {
        "status": "ok",
        "service": "Automatizador de Documentos AEP",
        "version": "0.2.0",
        "pipeline_ready": True,
        "processing": "temporary",
    }
    assert response.headers["cache-control"] == "no-store"
    assert response.headers["x-content-type-options"] == "nosniff"


def test_health_and_validation_fail_closed_without_verified_private_template(
    api_client: TestClient,
) -> None:
    pipeline = web_app.PIPELINE_INSTANCE
    pipeline.settings = replace(
        pipeline.settings,
        template_path=pipeline.settings.base_dir / "modelo-ausente.docx",
        template_manifest_path=(
            pipeline.settings.base_dir / "manifesto-ausente.json"
        ),
        allow_synthetic_template_fallback=False,
    )

    health = api_client.get("/api/health")
    validation = api_client.post("/api/validate")

    assert health.status_code == 503
    assert health.json()["status"] == "degraded"
    assert health.json()["pipeline_ready"] is False
    assert validation.status_code == 503
    assert validation.json()["detail"]["code"] == "pipeline_not_ready"


def test_home_page_loads_upload_and_generation_interface(
    api_client: TestClient,
) -> None:
    response = api_client.get("/")

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/html")
    assert "Automatizador de Documentos AEP" in response.text
    assert "Validar arquivos" in response.text
    assert "Gerar documento AEP" in response.text
    assert 'name="company_logo"' in response.text
    assert 'name="compatibility_mode"' in response.text
    assert "opcional" in response.text.casefold()

    script = api_client.get("/static/app.js")
    assert script.status_code == 200
    assert 'payload.set("compatibility_mode"' in script.text
    assert "function invalidateValidation()" in script.text


def test_multipart_validation_finds_synthetic_ghes_and_population(
    validate_job: dict[str, Any],
) -> None:
    assert validate_job["status"] == "needs_reconciliation"
    summary = validate_job["summary"]
    assert summary["ghe_count"] == 3
    assert summary["total_population"] == 16
    assert [ghe["code"] for ghe in summary["ghes"]] == [
        "GHE 10",
        "GHE 20",
        "GHE 30",
    ]
    assert [ghe["employees"] for ghe in summary["ghes"]] == [5, 8, 3]
    assert validate_job["errors"] == []


def test_validation_exposes_unmatched_ergo_block_for_reconciliation(
    validate_job: dict[str, Any],
) -> None:
    items = validate_job["reconciliation"]["items"]

    assert len(items) == 4
    assert sum(item["status"] == "auto_matched" for item in items) == 3
    pending = next(item for item in items if item["status"] == "needs_review")
    assert pending["source_code"] == "GHE 99"
    assert pending["candidates"] == []
    assert validate_job["warnings"]


def test_generation_poll_download_report_and_input_cleanup_without_logo(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    job_id = validate_job["job_id"]
    job_dir = web_app.JOBS[job_id].job_dir
    input_paths = tuple(web_app.JOBS[job_id].files.values())
    assert "company_logo" not in web_app.JOBS[job_id].files
    assert all(path.is_file() for path in input_paths)

    response = api_client.post(
        "/api/generate",
        json={
            "job_id": job_id,
            "reconciliations": [_extra_ergo_decision(validate_job)],
        },
    )
    assert response.status_code == 202, response.text
    assert response.json()["status"] == "generating"

    completed = _wait_for_terminal_status(api_client, job_id)
    assert completed["status"] == "completed", completed
    assert completed["progress"] == 100
    assert completed["downloads"]["document"]
    assert completed["downloads"]["validation_report"]

    document_response = api_client.get(
        f"/api/jobs/{job_id}/document"
    )
    assert document_response.status_code == 200
    assert document_response.content.startswith(b"PK\x03\x04")
    with zipfile.ZipFile(io.BytesIO(document_response.content)) as package:
        assert "word/document.xml" in package.namelist()
    generated = Document(io.BytesIO(document_response.content))
    visible_text = "\n".join(
        paragraph.text for paragraph in generated.paragraphs
    )
    assert "Empresa Sintética de Teste" in visible_text
    assert "GHE 10" in visible_text
    assert "Plano de ação integrado" in visible_text

    report_response = api_client.get(
        f"/api/jobs/{job_id}/validation-report"
    )
    assert report_response.status_code == 200
    report = json.loads(report_response.content)
    assert report["job_id"] == job_id
    assert report["status"] == "completed"
    assert report["summary"]["total_population"] == 16

    assert job_dir.is_dir()
    assert all(not path.exists() for path in input_paths)
    assert web_app.JOBS[job_id].document_path.is_file()
    assert job_id not in web_app.PIPELINE_INSTANCE._external_models


def test_mutating_synthetic_spreadsheet_changes_pipeline_and_docx(
    api_client: TestClient,
    valid_form_data: dict[str, str],
    valid_uploads: dict[str, tuple[str, bytes, str]],
) -> None:
    original_response = api_client.post(
        "/api/validate",
        data=valid_form_data,
        files=valid_uploads,
    )
    assert original_response.status_code == 200, original_response.text
    original_validation = original_response.json()
    original_document = _generate_and_download(
        api_client, original_validation
    )

    workbook = load_workbook(
        io.BytesIO(valid_uploads["ghe_spreadsheet"][1])
    )
    workbook.active["E2"] = 6
    mutated_spreadsheet = io.BytesIO()
    workbook.save(mutated_spreadsheet)
    workbook.close()

    mutated_uploads = dict(valid_uploads)
    mutated_uploads["ghe_spreadsheet"] = (
        "ghe_sinteticos_mutacao.xlsx",
        mutated_spreadsheet.getvalue(),
        valid_uploads["ghe_spreadsheet"][2],
    )
    mutated_response = api_client.post(
        "/api/validate",
        data=valid_form_data,
        files=mutated_uploads,
    )
    assert mutated_response.status_code == 200, mutated_response.text
    mutated_validation = mutated_response.json()
    assert original_validation["summary"]["total_population"] == 16
    assert mutated_validation["summary"]["total_population"] == 17
    assert mutated_validation["summary"]["ghes"][0]["employees"] == 6

    mutated_document = _generate_and_download(
        api_client, mutated_validation
    )
    assert _population_for_ghe(original_document, "GHE 10") == 5
    assert _population_for_ghe(mutated_document, "GHE 10") == 6


def test_separate_analysis_mode_validates_generates_and_reports(
    api_client: TestClient,
    valid_form_data: dict[str, str],
    valid_uploads: dict[str, tuple[str, bytes, str]],
) -> None:
    form_data = {**valid_form_data, "analysis_mode": "separate"}
    technical_bytes = valid_uploads["integrated_report"][1]
    technical_type = valid_uploads["integrated_report"][2]
    uploads = {
        key: value
        for key, value in valid_uploads.items()
        if key != "integrated_report"
    }
    uploads.update(
        {
            "psychosocial_analysis": (
                "agente_psicossocial_sintetico.docx",
                technical_bytes,
                technical_type,
            ),
            "ergonomic_analysis": (
                "agente_ergonomico_sintetico.docx",
                technical_bytes,
                technical_type,
            ),
        }
    )

    response = api_client.post(
        "/api/validate",
        data=form_data,
        files=uploads,
    )
    assert response.status_code == 200, response.text
    validation = response.json()
    record = web_app.JOBS[validation["job_id"]]
    assert "psychosocial_analysis" in record.files
    assert "ergonomic_analysis" in record.files
    assert "integrated_report" not in record.files

    document = _generate_and_download(api_client, validation)
    assert document.startswith(b"PK\x03\x04")
    report_response = api_client.get(
        f"/api/jobs/{validation['job_id']}/validation-report"
    )
    assert report_response.status_code == 200
    assert report_response.json()["document"]["analysis_mode"] == "separate"


def test_generation_requires_explicit_reconciliation(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    response = api_client.post(
        "/api/generate", json={"job_id": validate_job["job_id"]}
    )

    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "reconciliation_required"


def test_generation_rejects_compatibility_changed_after_validation(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    response = api_client.post(
        "/api/generate",
        json={
            "job_id": validate_job["job_id"],
            "reconciliations": [_extra_ergo_decision(validate_job)],
            "compatibility_mode": True,
        },
    )

    assert response.status_code == 409
    assert response.json()["detail"]["code"] == "validation_stale"
    record = web_app.JOBS[validate_job["job_id"]]
    assert record.compatibility is None
    assert record.fields["compatibility_mode"] is False


def test_compatibility_request_without_private_profile_fails_explicitly(
    api_client: TestClient,
    valid_form_data: dict[str, str],
    valid_uploads: dict[str, tuple[str, bytes, str]],
) -> None:
    form_data = {**valid_form_data, "compatibility_mode": "true"}

    response = api_client.post(
        "/api/validate", data=form_data, files=valid_uploads
    )

    assert response.status_code == 422
    detail = response.json()["detail"]
    assert detail["code"] == "compatibility_profile_unavailable"
    snapshot = api_client.get(f"/api/jobs/{detail['job_id']}").json()
    assert snapshot["compatibility"] is None
    report = api_client.get(
        snapshot["downloads"]["validation_report"]
    ).json()
    assert report["document"]["compatibility_mode"] is False


def test_private_profile_is_preserved_through_generation_and_reporting(
    api_client: TestClient,
    valid_form_data: dict[str, str],
    valid_uploads: dict[str, tuple[str, bytes, str]],
) -> None:
    pipeline = web_app.PIPELINE_INSTANCE
    private_dir = pipeline.settings.base_dir / "private_templates"
    private_dir.mkdir()
    profile_path = private_dir / "pilot-profile.json"
    source_paths = {
        "ghe_spreadsheet": (
            Path(__file__).resolve().parents[1]
            / "fixtures/public_synthetic/ghe_sinteticos.xlsx"
        ),
        "psychosocial_report": (
            Path(__file__).resolve().parents[1]
            / "fixtures/public_synthetic/psicossocial_sintetico.docx"
        ),
        "ergo_report": (
            Path(__file__).resolve().parents[1]
            / "fixtures/public_synthetic/ergo_sintetico.doc"
        ),
        "integrated_report": (
            Path(__file__).resolve().parents[1]
            / "fixtures/public_synthetic/tecnico_integrado_sintetico.docx"
        ),
        "cnpj_card": (
            Path(__file__).resolve().parents[1]
            / "fixtures/public_synthetic/cartao_cnpj_sintetico.png"
        ),
    }
    profile_path.write_text(
        json.dumps(
            {
                "schema_version": 1,
                "mode": "pilot_reference",
                "analysis_mode": "integrated",
                "input_fingerprint": _compatibility_input_fingerprint(
                    source_paths, "integrated"
                ),
                "included_ergo_ordinals": [1, 2, 3],
                "omitted_ergo_ordinals": [4],
            }
        ),
        encoding="utf-8",
    )
    pipeline.settings = replace(
        pipeline.settings, compatibility_profile_path=profile_path
    )
    response = api_client.post(
        "/api/validate",
        data={**valid_form_data, "compatibility_mode": "true"},
        files=valid_uploads,
    )
    assert response.status_code == 200, response.text
    validated = response.json()
    assert validated["compatibility"]["mode"] == "pilot_reference"

    generation = api_client.post(
        "/api/generate",
        json={
            "job_id": validated["job_id"],
            "reconciliations": [_extra_ergo_decision(validated)],
            "compatibility_mode": True,
        },
    )
    assert generation.status_code == 202, generation.text
    completed = _wait_for_terminal_status(
        api_client, validated["job_id"]
    )
    assert completed["status"] == "completed", completed
    assert completed["compatibility"]["mode"] == "pilot_reference"
    report = api_client.get(
        completed["downloads"]["validation_report"]
    ).json()
    assert report["document"]["compatibility_mode"] is True
    assert report["compatibility"]["omitted_ergo_source_ids"]


def test_rejects_extension_with_incompatible_real_type(
    api_client: TestClient,
    valid_form_data: dict[str, str],
    valid_uploads: dict[str, tuple[str, bytes, str]],
) -> None:
    invalid_uploads = dict(valid_uploads)
    invalid_uploads["ghe_spreadsheet"] = (
        "ghe_sinteticos.xlsx",
        b"conteudo sintetico que nao e um pacote xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    response = api_client.post(
        "/api/validate", data=valid_form_data, files=invalid_uploads
    )

    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "type_mismatch"
    assert web_app.JOBS == {}
    assert list(web_app.RUNTIME_ROOT.iterdir()) == []


def test_rejects_macro_before_pipeline_extractors(
    api_client: TestClient,
    valid_form_data: dict[str, str],
    valid_uploads: dict[str, tuple[str, bytes, str]],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    original = valid_uploads["integrated_report"]
    source = io.BytesIO(original[1])
    mutated = io.BytesIO()
    with zipfile.ZipFile(source) as incoming, zipfile.ZipFile(
        mutated, "w", compression=zipfile.ZIP_DEFLATED
    ) as outgoing:
        for item in incoming.infolist():
            outgoing.writestr(item, incoming.read(item.filename))
        outgoing.writestr("word/vbaProject.bin", b"macro sintetica sem validade")
    invalid_uploads = dict(valid_uploads)
    invalid_uploads["integrated_report"] = (
        original[0],
        mutated.getvalue(),
        original[2],
    )

    def extractor_must_not_run(*args: Any, **kwargs: Any) -> None:
        raise AssertionError("a pipeline não deve receber Office inseguro")

    monkeypatch.setattr(
        web_app.PIPELINE_INSTANCE,
        "validate",
        extractor_must_not_run,
    )
    response = api_client.post(
        "/api/validate",
        data=valid_form_data,
        files=invalid_uploads,
    )

    assert response.status_code == 422
    assert response.json()["detail"]["code"] == "type_mismatch"
    assert web_app.JOBS == {}


def test_stream_limit_does_not_trust_content_length(
    api_client: TestClient,
    valid_form_data: dict[str, str],
    valid_uploads: dict[str, tuple[str, bytes, str]],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(web_app, "MAX_REQUEST_SIZE", 1024)

    response = api_client.post(
        "/api/validate",
        data=valid_form_data,
        files=valid_uploads,
        headers={"Content-Length": "1"},
    )

    assert response.status_code == 413
    assert response.json()["detail"]["code"] == "request_too_large"
    assert web_app.JOBS == {}


def test_multipart_rejects_too_many_file_parts(
    api_client: TestClient,
    valid_form_data: dict[str, str],
    valid_uploads: dict[str, tuple[str, bytes, str]],
) -> None:
    png = valid_uploads["cnpj_card"]
    files = list(valid_uploads.items())
    files.extend(
        (f"extra_{index}", (f"extra-{index}.png", png[1], png[2]))
        for index in range(web_app.MAX_UPLOAD_FILES - len(files) + 1)
    )

    response = api_client.post(
        "/api/validate",
        data=valid_form_data,
        files=files,
    )

    assert response.status_code == 413
    assert response.json()["detail"]["code"] == "multipart_limit"
    assert web_app.JOBS == {}


def test_expired_job_cleanup_removes_only_its_runtime_directory(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    job_id = validate_job["job_id"]
    record = web_app.JOBS[job_id]
    unrelated = web_app.RUNTIME_ROOT.parent / "nao-remover.txt"
    unrelated.write_text("fixture sintética", encoding="utf-8")
    record.updated_at -= web_app.JOB_TTL_SECONDS + 1

    removed = web_app.cleanup_expired_jobs(now=time.time())

    assert removed == 1
    assert job_id not in web_app.JOBS
    assert not record.job_dir.exists()
    assert unrelated.read_text(encoding="utf-8") == "fixture sintética"


def test_orphan_cleanup_retries_after_transient_failure(
    api_client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    orphan = web_app.RUNTIME_ROOT / ("a" * 32)
    orphan.mkdir()
    old_timestamp = time.time() - web_app.JOB_TTL_SECONDS - 1
    os.utime(orphan, (old_timestamp, old_timestamp))
    original_rmtree = shutil.rmtree
    attempts = 0

    def blocked_rmtree(path: Path, *, ignore_errors: bool = False) -> None:
        nonlocal attempts
        if Path(path).resolve() == orphan.resolve():
            attempts += 1
            raise PermissionError("fixture sintética em uso")
        original_rmtree(path, ignore_errors=ignore_errors)

    monkeypatch.setattr(shutil, "rmtree", blocked_rmtree)
    first = web_app.cleanup_expired_jobs(now=time.time())

    assert first == 0
    assert attempts == 3
    assert orphan.is_dir()

    monkeypatch.setattr(shutil, "rmtree", original_rmtree)
    second = web_app.cleanup_expired_jobs(now=time.time())

    assert second == 1
    assert not orphan.exists()


def test_cors_allows_only_configured_github_pages_origin(
    api_client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(web_app, "REQUIRE_ORIGIN", True)

    allowed = api_client.get(
        "/api/jobs/" + ("a" * 32),
        headers={"Origin": "https://kinhoog.github.io"},
    )
    rejected = api_client.get(
        "/api/jobs/" + ("a" * 32),
        headers={"Origin": "https://example.invalid"},
    )
    missing = api_client.get("/api/jobs/" + ("a" * 32))
    preflight = api_client.options(
        "/api/generate",
        headers={
            "Origin": "https://kinhoog.github.io",
            "Access-Control-Request-Method": "POST",
            "Access-Control-Request-Headers": "content-type",
        },
    )

    assert allowed.status_code == 404
    assert (
        allowed.headers["access-control-allow-origin"]
        == "https://kinhoog.github.io"
    )
    exposed_headers = allowed.headers["access-control-expose-headers"].casefold()
    assert "x-aep-content-length" in exposed_headers
    assert "x-aep-content-sha256" in exposed_headers
    assert rejected.status_code == 403
    assert "access-control-allow-origin" not in rejected.headers
    assert missing.status_code == 403
    assert preflight.status_code == 200
    assert (
        preflight.headers["access-control-allow-origin"]
        == "https://kinhoog.github.io"
    )
    assert "*" not in preflight.headers["access-control-allow-methods"]


def test_every_api_response_disables_caching(
    api_client: TestClient,
) -> None:
    response = api_client.get("/api/jobs/" + ("f" * 32))

    assert response.status_code == 404
    assert response.headers["cache-control"] == "no-store"
    assert response.headers["pragma"] == "no-cache"


def test_complete_document_download_then_explicit_idempotent_deletion(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    job_id = validate_job["job_id"]
    document = _generate_and_download(api_client, validate_job)
    job_dir = web_app.JOBS[job_id].job_dir

    assert document.startswith(b"PK\x03\x04")
    with zipfile.ZipFile(io.BytesIO(document)) as package:
        assert "word/document.xml" in package.namelist()
    assert job_id in web_app.JOBS
    assert job_dir.is_dir()

    first = api_client.delete(f"/api/jobs/{job_id}")
    second = api_client.delete(f"/api/jobs/{job_id}")

    assert first.status_code == 204
    assert second.status_code == 204
    assert job_id not in web_app.JOBS
    assert not job_dir.exists()


def test_delete_returns_conflict_while_download_is_active(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    job_id = validate_job["job_id"]
    record = web_app.JOBS[job_id]
    record.active_downloads = 1

    busy = api_client.delete(f"/api/jobs/{job_id}")

    assert busy.status_code == 409
    assert busy.json()["detail"]["code"] == "job_busy"
    assert web_app.JOBS[job_id] is record
    assert record.job_dir.is_dir()

    record.active_downloads = 0
    assert api_client.delete(f"/api/jobs/{job_id}").status_code == 204
    assert job_id not in web_app.JOBS
    assert not record.job_dir.exists()


def test_download_endpoint_deletes_only_after_complete_response(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    job_id = validate_job["job_id"]
    generation = api_client.post(
        "/api/generate",
        json={
            "job_id": job_id,
            "reconciliations": [_extra_ergo_decision(validate_job)],
        },
    )
    assert generation.status_code == 202
    completed = _wait_for_terminal_status(api_client, job_id)
    assert completed["downloads"]["download"].endswith("/download")
    job_dir = web_app.JOBS[job_id].job_dir

    partial = api_client.get(
        completed["downloads"]["download"],
        headers={"Range": "bytes=0-3"},
    )

    assert partial.status_code == 416
    assert partial.json()["detail"]["code"] == "partial_download_not_supported"
    assert job_id in web_app.JOBS
    assert job_dir.is_dir()

    response = api_client.get(completed["downloads"]["download"])

    assert response.status_code == 200
    assert response.content.startswith(b"PK\x03\x04")
    assert int(response.headers["content-length"]) == len(response.content)
    assert int(response.headers["x-aep-content-length"]) == len(response.content)
    assert response.headers["x-aep-content-sha256"] == hashlib.sha256(
        response.content
    ).hexdigest()
    assert response.headers["cache-control"] == "no-store, no-transform"
    with zipfile.ZipFile(io.BytesIO(response.content)) as package:
        assert package.testzip() is None
    assert job_id not in web_app.JOBS
    assert not job_dir.exists()
    assert api_client.get(f"/api/jobs/{job_id}").status_code == 404


def test_abandoned_validated_job_is_removed_by_ttl(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    job_id = validate_job["job_id"]
    record = web_app.JOBS[job_id]
    record.updated_at = time.time() - web_app.JOB_TTL_SECONDS - 1

    assert web_app.cleanup_expired_jobs() == 1
    assert job_id not in web_app.JOBS
    assert not record.job_dir.exists()


def test_ttl_cleanup_does_not_remove_an_active_download(
    api_client: TestClient,
    validate_job: dict[str, Any],
) -> None:
    job_id = validate_job["job_id"]
    record = web_app.JOBS[job_id]
    record.updated_at = time.time() - web_app.JOB_TTL_SECONDS - 1
    record.active_downloads = 1

    assert web_app.cleanup_expired_jobs() == 0
    assert web_app.JOBS[job_id] is record
    assert record.job_dir.is_dir()

    record.active_downloads = 0
    assert web_app.cleanup_expired_jobs() == 1
    assert job_id not in web_app.JOBS
    assert not record.job_dir.exists()
